#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
منصة التدقيق الديناميكي الموحد V3.0 — وثيقة شرح الواجهات والحسابات
يولّد وثيقة Word شاملة بالعربية (RTL) تشرح كل واجهة وكل محرك حسابي.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ═══════════════════════════════════════════════════════════════════════
# إعدادات عامة
# ═══════════════════════════════════════════════════════════════════════

OUTPUT_PATH = "/home/z/my-project/download/منصة_التدقيق_الديناميكي_V3_وثيقة_الواجهات_والحسابات.docx"

# ألوان
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0xC9, 0xA2, 0x27)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF5, 0xF7, 0xFA)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xB7, 0x71, 0x0F)


def set_rtl(paragraph):
    """جعل الفقرة من اليمين إلى اليسار"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def set_rtl_run(run):
    """جعل النص داخل الـ run من اليمين لليسار"""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def set_cell_shading(cell, color_hex):
    """تلوين خلفية الخلية"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_table_rtl(table):
    """جعل الجدول RTL"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)


def add_heading_ar(doc, text, level=1, color=NAVY, size=None):
    """إضافة عنوان عربي"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.keep_with_next = True

    run = p.add_run(text)
    run.font.name = 'Calibri'
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), 'Calibri')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = size or Pt({1: 20, 2: 16, 3: 13, 4: 11}.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = color
    set_rtl_run(run)
    return p


def add_para_ar(doc, text, bold=False, italic=False, color=DARK,
                size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, indent_first=False):
    """إضافة فقرة عربية"""
    p = doc.add_paragraph()
    p.alignment = align
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(6)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.6)

    run = p.add_run(text)
    run.font.name = 'Calibri'
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), 'Calibri')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    set_rtl_run(run)
    return p


def add_bullet_ar(doc, text, level=0, bold=False):
    """نقطة قائمة عربية"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.right_indent = Cm(0.75 * (level + 1))

    # علامة النقطة (نستخدم رمز من اليمين)
    bullets = ['◀', '◇', '•']
    marker = bullets[min(level, 2)]
    run0 = p.add_run(f'{marker}  ')
    run0.font.name = 'Calibri'
    run0.font.size = Pt(11)
    run0.font.color.rgb = NAVY
    run0.font.bold = True
    set_rtl_run(run0)

    run = p.add_run(text)
    run.font.name = 'Calibri'
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), 'Calibri')
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.color.rgb = DARK
    set_rtl_run(run)
    return p


def add_formula_box(doc, formula_text, label=""):
    """صندوق معادلة مميز"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)

    # خلفية صفراء فاتحة
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FBF7E8')
    pPr.append(shd)

    # حدود
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '6')
        b.set(qn('w:color'), 'C9A227')
        pBdr.append(b)
    pPr.append(pBdr)

    if label:
        run_lbl = p.add_run(f'{label}:  ')
        run_lbl.font.name = 'Consolas'
        run_lbl.font.size = Pt(11)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = NAVY

    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    return p


def add_table_simple(doc, headers, rows, header_bg='1F3A5F',
                     col_widths=None, font_size=10):
    """جدول بسيط RTL"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)

    # رؤوس
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = WHITE
        set_rtl_run(run)
        set_cell_shading(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # بيانات
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(font_size)
            run.font.color.rgb = DARK
            set_rtl_run(run)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cells[c_idx], 'F5F7FA')

    # عرض الأعمدة
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_info_box(doc, title, content, color_bg='E8F0FE', color_border='1F3A5F',
                 icon='ⓘ'):
    """صندوق معلومات"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, color_bg)

    # حدود ملوّنة
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:color'), color_border)
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # عنوان
    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_title)
    run_t = p_title.add_run(f'{icon}  {title}')
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(12)
    run_t.font.bold = True
    run_t.font.color.rgb = NAVY
    set_rtl_run(run_t)

    # محتوى
    if isinstance(content, str):
        content = [content]
    for line in content:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
        set_rtl_run(run)

    doc.add_paragraph()  # فاصل
    return table


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ═══════════════════════════════════════════════════════════════════════
# بدء الوثيقة
# ═══════════════════════════════════════════════════════════════════════

doc = Document()

# إعداد الخط الافتراضي
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:cs'), 'Calibri')
rFonts.set(qn('w:ascii'), 'Calibri')
rFonts.set(qn('w:hAnsi'), 'Calibri')

# إعداد هوامش الصفحة
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# جعل القسم RTL
sectPr = section._sectPr
bidi = OxmlElement('w:bidi')
sectPr.append(bidi)

print("✓ تم تهيئة الوثيقة")
print("✓ بدء كتابة المحتوى...")

# ═══════════════════════════════════════════════════════════════════════
# الغلاف
# ═══════════════════════════════════════════════════════════════════════

# شعار علوي
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('◆ ◆ ◆')
run.font.size = Pt(28)
run.font.color.rgb = ACCENT
run.font.bold = True

# العنوان الرئيسي
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(12)
run = p.add_run('منصة التدقيق الديناميكي الموحد')
run.font.name = 'Calibri'
rPr = run._r.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:cs'), 'Calibri')
rPr.append(rFonts)
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = NAVY
set_rtl_run(run)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run('V3.0')
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = ACCENT

# خط فاصل
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━')
run.font.size = Pt(12)
run.font.color.rgb = ACCENT

# العنوان الفرعي
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run('وثيقة شرح الواجهات والحسابات التفصيلية')
run.font.name = 'Calibri'
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = DARK
set_rtl_run(run)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)
run = p.add_run('المدخلات • المخرجات • المعادلات الحسابية • العلاقات المعتمدة')
run.font.name = 'Calibri'
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = GRAY
set_rtl_run(run)

# بطاقة المعلومات
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(info_table)
info_data = [
    ('اسم المنصة', 'منصة التدقيق الديناميكي الموحد V3.0'),
    ('الإصدار', '3.0 — الإصدار الإنتاجي'),
    ('المرجع العلمي', 'الأطروحة + ملفات Excel I-1 إلى I-9'),
    ('الكود الحاكم', 'الكود السوري 2024 + UFC 3-340-02'),
    ('تاريخ الإصدار', '2026-06-23'),
]
for i, (k, v) in enumerate(info_data):
    c1 = info_table.rows[i].cells[0]
    c2 = info_table.rows[i].cells[1]
    set_cell_shading(c1, '1F3A5F')
    set_cell_shading(c2, 'F5F7FA')
    c1.text = ''
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p1)
    r1 = p1.add_run(k)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = WHITE
    set_rtl_run(r1)
    c2.text = ''
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p2)
    r2 = p2.add_run(v)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    set_rtl_run(r2)
    c1.width = Cm(5)
    c2.width = Cm(10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════
# جدول المحتويات (يدوي)
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'جدول المحتويات', level=1, color=NAVY)

toc_items = [
    ('الفصل الأول: مقدمة عامة عن المنصة', '4'),
    ('الفصل الثاني: البنية المعمارية وطبقات الحساب', '6'),
    ('الفصل الثالث: خط الحساب الموحد (Pipeline)', '8'),
    ('الفصل الرابع: واجهات المصادقة وإدارة المستخدمين', '10'),
    ('    4.1 صفحة تسجيل الدخول /auth/login', '10'),
    ('    4.2 صفحة التسجيل /auth/register', '12'),
    ('    4.3 صفحة انتظار الموافقة /auth/pending', '13'),
    ('    4.4 لوحة المدير /admin', '14'),
    ('الفصل الخامس: محرك الاختراق (Penetration Engine)', '16'),
    ('الفصل السادس: محرك الضغط العصفي (Blast Pressure)', '19'),
    ('الفصل السابع: محرك التصميم الإنشائي (Structural Concrete)', '22'),
    ('الفصل الثامن: محرك تصميم التسليح (Rebar Design)', '25'),
    ('الفصل التاسع: محرك المفاضلة الهندسية (Geometry Comparison)', '28'),
    ('الفصل العاشر: صفحات لوحة التحكم التفصيلية', '30'),
    ('    10.1 الصفحة الرئيسية /dashboard', '30'),
    ('    10.2 صفحة الاختراق /dashboard/penetration', '33'),
    ('    10.3 صفحة أحمال الانفجار /dashboard/blast-loads', '35'),
    ('    10.4 صفحة التصميم الإنشائي /dashboard/structural-design', '37'),
    ('    10.5 صفحة التسليح /dashboard/rebar-design', '39'),
    ('    10.6 صفحة المفاضلة /dashboard/comparison', '41'),
    ('    10.7 صفحة خط الأساس /dashboard/benchmark', '43'),
    ('    10.8 صفحة المتغيرات الموحدة /dashboard/variables', '45'),
    ('    10.9 صفحة مصفوفة التتبع /dashboard/rtm', '47'),
    ('    10.10 صفحة التقارير /dashboard/reports', '49'),
    ('    10.11 صفحة الإعدادات /dashboard/settings', '50'),
    ('    10.12 صفحة عن المنصة /dashboard/about', '51'),
    ('الفصل الحادي عشر: المعادلات الحسابية المرجعية', '52'),
    ('الفصل الثاني عشر: جدول المتغيرات الموحدة (42 متغير)', '55'),
    ('الفصل الثالث عشر: خلاصة وتوصيات', '57'),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    # عنوان
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = title.strip().startswith('الفصل')
    run.font.color.rgb = NAVY if title.strip().startswith('الفصل') else DARK
    set_rtl_run(run)
    # نقاط
    dots_run = p.add_run('  ' + '·' * 50 + '  ')
    dots_run.font.size = Pt(9)
    dots_run.font.color.rgb = GRAY
    # رقم الصفحة
    page_run = p.add_run(page)
    page_run.font.name = 'Calibri'
    page_run.font.size = Pt(11)
    page_run.font.bold = True
    page_run.font.color.rgb = ACCENT

add_page_break(doc)

print("✓ تم كتابة الغلاف وجدول المحتويات")

# ═══════════════════════════════════════════════════════════════════════
# الفصل الأول: مقدمة عامة عن المنصة
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الأول: مقدمة عامة عن المنصة', level=1, color=NAVY)

add_heading_ar(doc, '1.1 الهدف من المنصة', level=2)
add_para_ar(doc,
    'منصة التدقيق الديناميكي الموحد V3.0 هي منصة هندسية متخصصة في تدقيق وحساب '
    'المنشآت الوقائية المعرضة لأحمال الانفجار والاختراق. تجمع المنصة بين محركات '
    'حسابية متخصصة (محرك الاختراق، محرك الضغط العصفي، محرك التصميم الإنشائي، '
    'ومحرك المفاضلة الهندسية) في خط أنابيب (Pipeline) موحد يأخذ مدخلات المستخدم '
    'وينتج توصية هندسية نهائية بشأن الشكل الإنشائي الأمثل للمقطع (مستطيل / دائري / مقوّس).',
    indent_first=True)

add_para_ar(doc,
    'تعتمد المنصة على معادلات مرجعية مستخرجة من أطروحة علمية متخصصة، ومُنَضَّدة '
    'بملفات Excel (I-1 إلى I-9) لجداول الاستيفاء. كما تستند إلى الكود السوري 2024 '
    'للتصميم بالطريقة المرنة (WSD) وإلى دليل UFC 3-340-02 الأمريكي للمنشآت الحمائية '
    'ضد الأحمال الديناميكية. كل معادلة في المنصة معرّفة كدالة نقية (Pure Function) '
    'قابلة للاختبار المستقل، مما يضمن قابلية التتبع والتحقق من كل خطوة حسابية.',
    indent_first=True)

add_heading_ar(doc, '1.2 المكونات الرئيسية', level=2)
add_para_ar(doc, 'تتكوّن المنصة من أربعة محركات حسابية رئيسية، ينسّق بينها منسّق عام (Orchestrator):')

add_table_simple(doc,
    ['#', 'المحرك', 'الوظيفة', 'الملف المرجعي'],
    [
        ['1', 'محرك الاختراق', 'حساب عمق اختراق القنبلة في التربة/الخرسانة',
         'penetration-core.ts'],
        ['2', 'محرك الضغط العصفي', 'حساب ضغط الانفجار والمدة الزمنية للموجة',
         'blast-pressure-core.ts + blast-loads.ts'],
        ['3', 'محرك التصميم الإنشائي', 'حساب سماكة المقطع والتسليح المطلوب والتحقق',
         'structural-concrete-core.ts'],
        ['4', 'محرك تصميم التسليح', 'تفصيل التسليح: عدد القضبان والقطر والفحوصات',
         'rebar.ts'],
        ['5', 'محرك المفاضلة', 'مقارنة الأشكال الثلاثة واختيار الأمثل',
         'geometry-comparator.ts'],
        ['—', 'المنسّق العام (Orchestrator)', 'يربط المحركات بتسلسل صحيح',
         'orchestrator.ts'],
    ],
    col_widths=[1, 4, 8, 4])

add_heading_ar(doc, '1.3 الجمهور المستهدف', level=2)
add_para_ar(doc,
    'تستهدف المنصة المهندسين الإنشائيين العاملين في تصميم المنشآت الحمائية، '
    'والباحثين الأكاديميين في مجال ديناميكا الانفجارات، والمقدّرين الفنيين '
    'للأضرار، والجهات السيادية المسؤولة عن البنية التحتية الحساسة. كما تتيح '
    'المنصة للمدير الحوكمي إدارة المستخدمين والصلاحيات عبر لوحة تحكم مخصصة، '
    'مع تسجيل كامل لسجل التدقيق (Audit Trail) لكل عملية حسابية.',
    indent_first=True)

add_heading_ar(doc, '1.4 الميزات الجوهرية', level=2)
features = [
    ('التحقق التلقائي (Validation)', 'كل مدخل يمر عبر طبقة Zod للتحقق من النطاق والنوع قبل الدخول للمحرك، مع رسائل خطأ عربية واضحة.'),
    ('القيم المقفلة (Locked Values)', 'القيم التي يُنتجها محرك ويستهلكها المحرك التالي تكون مقفلة (للقراءة فقط) لمنع الكتابة المتكررة والانحراف عن المرجع.'),
    ('مقارنة BMK-02 المرجعية', 'كل نتيجة من المحركات تُقارن آلياً مع الحالة المرجعية BMK-02 (قنبلة MK83 في تربة متوسطة) لحساب نسبة الانحراف.'),
    ('42 متغير موحد', 'نموذج متغيرات موحّد يصنّف كل قيمة في المنصة إلى 5 فئات (input/lookup/computed/locked/output) و3 مسارات (سقف/جدار/مشترك).'),
    ('دعم Offline/PWA', 'المنصة تعمل كتطبيق ويب تقدمي (PWA) مع IndexedDB للتخزين المحلي ومزامنة لاحقة عند توفر الاتصال.'),
    ('الأمان والحوكمة', 'تشفير AES-256-GCM لملفات تعريف الجلسة + bcrypt لكلمات المرور + RBAC بـ 7 أذونات تفصيلية + حالة اشتراك PENDING حتى موافقة المدير.'),
]
for title, desc in features:
    add_bullet_ar(doc, f'{title}: {desc}', bold=False)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════
# الفصل الثاني: البنية المعمارية وطبقات الحساب
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الثاني: البنية المعمارية وطبقات الحساب', level=1, color=NAVY)

add_heading_ar(doc, '2.1 البنية التقنية العامة', level=2)
add_para_ar(doc,
    'المنصة مبنية على إطار العمل Next.js 16 (App Router) مع TypeScript بصيغة صارمة، '
    'وتستخدم Tailwind CSS 4 لإدارة التنسيقات ومكتبة shadcn/ui لمكوّنات الواجهة. '
    'تُدار الحالة على العميل عبر React hooks وContext API، بينما تُنفَّذ العمليات الحسابية '
    'في طبقة Server Actions وAPI Routes لتقليل حمولة المتصفح وضمان سلامة البيانات.',
    indent_first=True)

add_para_ar(doc,
    'تتصل المنصة بقاعدة بيانات Supabase (PostgreSQL) عبر REST API (PostgREST) '
    'باستخدام محوّل مخصص (supabase-adapter.ts) يترجم استدعاءات Prisma التقليدية '
    'إلى طلبات REST. هذا الخيار يتيح نشر المنصة على Netlify بدون الحاجة إلى '
    'اتصال TCP مباشر بقاعدة البيانات، وهو مناسب لبيئات Serverless.',
    indent_first=True)

add_heading_ar(doc, '2.2 طبقات الحساب السبع', level=2)
add_para_ar(doc,
    'تتبع المنصة معمارية متعددة الطبقات (Layered Architecture) تفصل المسؤوليات بدقة. '
    'كل طبقة لها واجهة عقد (Contract) واضحة مع الطبقة المجاورة، ولا تُتجاوز طبقة '
    'إلى أخرى مباشرة. الجدول التالي يلخّص هذه الطبقات من الأعلى إلى الأسفل:')

add_table_simple(doc,
    ['الطبقة', 'المسؤولية', 'الملفات الرئيسية'],
    [
        ['1. واجهة المستخدم (UI)', 'عرض النتائج وقبول مدخلات المستخدم',
         'src/app/dashboard/*/page.tsx + src/components/*'],
        ['2. إجراءات العميل (Client Actions)', 'تحويل النموذج لطلب API وإدارة الحالة',
         'src/components/v3-engine-form.tsx + src/components/results-panel.tsx'],
        ['3. طبقة API (REST Routes)', 'استقبال الطلبات وتحقق الجلسة',
         'src/app/api/*/route.ts'],
        ['4. إجراءات الخادم (Server Actions)', 'تنفيذ العمليات الحساسة مع قاعدة البيانات',
         'src/app/actions/auth.actions.ts + admin.actions.ts'],
        ['5. المنسّق (Orchestrator)', 'تنسيق المحركات الأربعة بتسلسل صحيح',
         'src/lib/engine/orchestrator.ts'],
        ['6. المحركات (Engines)', 'الحسابات الفعلية لكل مرحلة',
         'src/lib/engine/*-core.ts'],
        ['7. الثوابت (Constants)', 'معاملات المرجعية والقيم المقفلة',
         'src/lib/engine/constants/* + src/lib/constants/reference-data.ts'],
    ],
    col_widths=[3.5, 6, 7])

add_heading_ar(doc, '2.3 نموذج المتغيرات الموحد', level=2)
add_para_ar(doc,
    'تتبنى المنصة نموذج متغيرات موحّد يصنّف كل قيمة في النظام إلى إحدى خمس فئات، '
    'مما يمنع التضارب بين المحركات ويضمن قابلية التتبع. هذا التصنيف هو العقد الذي '
    'يحكم تدفق البيانات بين الطبقات:')

add_table_simple(doc,
    ['الفئة', 'الرمز', 'الوصف', 'مثال'],
    [
        ['مدخل', 'input', 'قيمة يدخلها المستخدم مباشرة في النموذج', 'وزن القنبلة P=441kg'],
        ['مرجعي', 'lookup', 'قيمة من جداول الإكسيل أو الكود', 'K1=1.639 لـ Tritonal 80/20'],
        ['محسوب', 'computed', 'يُحسب داخل محرك ولا ينتقل لآخر', 'lambda1, lambda2'],
        ['مقفل', 'locked', 'يُنتج من محرك ويُستهلك في التالي', 'C_ef, h_pr, R_actual'],
        ['مخرج', 'output', 'النتيجة النهائية لخط الأنابيب', 'Pp, As, h_required'],
    ],
    col_widths=[2.5, 2, 7, 5])

add_heading_ar(doc, '2.4 فصل المسارات (Path Separation)', level=2)
add_para_ar(doc,
    'بعد الخطوة 4 (اعتماد السماكات الأولية)، ينقسم خط الحساب إلى مسارين مستقلين: '
    'مسار السقف (Roof Path) ومسار الجدار (Wall Path). كل مسار يحتسب أحماله '
    'وتصميمه الإنشائي باستقلال تام، مع اشتراكهما في القيم المقفلة من الخطوات 2-4. '
    'هذا الفصل ينعكس في بنية المخرجات: BlastLoadOutput و SectionDesignResult '
    'لها نسختان منفصلتان لكل مسار.',
    indent_first=True)

add_info_box(doc, 'ملاحظة هامة',
    ['القيم المقفلة (locked) من الخطوات 2-4 (C_ef, h_pr, R_actual, ht, Bt) تتدفق إلى كلا المسارين دون تعديل. هذا الضمان يمنع الانحراف التراكمي عن الحالة المرجعية BMK-02.'],
    color_bg='FFF8E1', color_border='C9A227', icon='⚠')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════
# الفصل الثالث: خط الحساب الموحد
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الثالث: خط الحساب الموحد (Pipeline)', level=1, color=NAVY)

add_heading_ar(doc, '3.1 نظرة شاملة على خط الحساب', level=2)
add_para_ar(doc,
    'خط الحساب في المنصة يتكوّن من 7 خطوات متسلسلة (مرقّمة 2 إلى 8 في المرجع الأصلي)، '
    'تبدأ من مدخلات المستخدم وتنتهي بقرار التصميم النهائي. كل خطوة تنتج قيماً '
    'تُستخدم كمدخلات للخطوة التالية، مع الحفاظ على عزل المسؤوليات بين المحركات. '
    'الشكل التالي يوضح تسلسل الحساب:')

# مخطط نصي لخط الأنابيب
add_info_box(doc, 'مخطط خط الحساب (Pipeline)',
    [
        'الخطوة 2: المدخلات والمعاملات (Inputs + Lookups)',
        '         ↓',
        'الخطوة 3: محرك الاختراق (Penetration) → C_ef, h_pr, R_actual',
        '         ↓',
        'الخطوة 4: اعتماد السماكات الأولية (Locked) → ht, Bt',
        '         ↓                                     ↓',
        'الخطوة 5: الأحمال (سقف)         |    الخطوة 5: الأحمال (جدار)',
        '         ↓                                     ↓',
        'الخطوة 6: العزم والتسليح (سقف)  |    الخطوة 6: العزم والتسليح (جدار)',
        '         ↓                                     ↓',
        'الخطوة 7: تصميم السقف النهائي   |    الخطوة 8: تصميم الجدار النهائي',
        '         ↓                                     ↓',
        '         └─────────────← المقارنة ←─────────────┘',
        '                         ↓',
        '                  القرار النهائي + التوصية',
    ],
    color_bg='E8F0FE', color_border='1F3A5F', icon='◆')

add_heading_ar(doc, '3.2 جدول الخطوات التفصيلي', level=2)

add_table_simple(doc,
    ['الخطوة', 'الاسم', 'المحرك المسؤول', 'المدخلات', 'المخرجات الرئيسية'],
    [
        ['2', 'المدخلات والمعاملات', '—',
         'السلاح + التربة + الهندسة', 'K1, kpr, gamma_b, gamma_g'],
        ['3', 'الاختراق', 'penetration-core.ts',
         'weaponId, impactVelocity, soilTypeCode',
         'C_ef, h_pr, R_actual, h_z, Zp'],
        ['4', 'السماكات المقفلة', 'reference-data.ts',
         'Hp, Hc, Hf من المرجع',
         'ht=107.22cm, Bt=8.05m, Hpc=3.83m'],
        ['5', 'الأحمال الانفجارية', 'blast-loads.ts',
         'C_ef, R_actual, Z, Hp_cm',
         'Pmax, P_ekv, Pct, Pp, tau, omega'],
        ['6', 'العزم الديناميكي', 'structural.ts',
         'Pp, ap, bp, lambda',
         'Mp (العزم البلاستيكي)'],
        ['7', 'تصميم السقف', 'structural-concrete-core.ts',
         'Mp, h0, Rbd, Rsd',
         'Hp_final, As, rho, mu'],
        ['8', 'تصميم الجدار', 'rebar.ts',
         'Mp, h0, Rbd, Rsd',
         'Hc_final, As, rho, bar_count'],
    ],
    col_widths=[1.2, 3, 3.5, 3.5, 5], font_size=9)

add_heading_ar(doc, '3.3 عقد المنسّق (Orchestrator Contract)', level=2)
add_para_ar(doc,
    'المنسّق العام هو نقطة الدخول الوحيدة لخط الحساب. يستقبل كائن EngineInput '
    'موحد ويُعيد كائن EngineOutput شامل. لا يحتوي المنسّق على أي منطق حسابي — '
    'مهمته الوحيدة هي ترتيب الاستدعاءات وجمع التحذيرات والأخطاء. هذا الفصل '
    'يضمن أن المعادلات مرجعية فقط للمحركات المتخصصة، ويمكن اختبار كل محرك '
    'بشكل مستقل.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات الموحدة (EngineInput)', level=3)
add_table_simple(doc,
    ['الحقل', 'النوع', 'الوحدة', 'الوصف'],
    [
        ['penetration.weaponId', 'string', '—', 'معرف السلاح من المكتبة'],
        ['penetration.impactVelocity', 'number', 'm/s', 'سرعة اصطدام القنبلة'],
        ['penetration.soilTypeCode', 'enum', '—', 'نوع التربة (6 خيارات)'],
        ['penetration.impactAngleDeg', 'number?', 'درجة', 'زاوية الاصطدام (0=عمودي)'],
        ['blast.radialDistance', 'number', 'm', 'البعد الشعاعي عن مركز الانفجار'],
        ['blast.ceilingDepth', 'number?', 'm', 'عمق السقف فوق مركز الانفجار'],
        ['design.tunnelSpanShort', 'number', 'm', 'البحر القصير للنفق (ap)'],
        ['design.tunnelSpanLong', 'number', 'm', 'البحر الطويل للنفق (bp)'],
        ['design.fcMpa', 'number', 'MPa', 'مقاومة ضغط الخرسانة'],
        ['design.fyMpa', 'number', 'MPa', 'إجهاد خضوع الحديد'],
        ['design.slabThicknessHintMm', 'number?', 'mm', 'سماكة مقترحة (اختياري)'],
        ['design.reinforcementRatio', 'number?', '—', 'نسبة تسليح مقترحة'],
    ],
    col_widths=[5, 2, 2, 7], font_size=10)

add_heading_ar(doc, 'المخرجات الموحدة (EngineOutput)', level=3)
add_table_simple(doc,
    ['الحقل', 'النوع', 'الوصف'],
    [
        ['inputs', 'EngineInput', 'المدخلات كما استُلمت (للتدقيق)'],
        ['intermediates.penetration', 'PenetrationOutput', 'كل مخرجات محرك الاختراق'],
        ['intermediates.blast', 'BlastOutput', 'كل مخرجات محرك الضغط'],
        ['structural', 'Record<Geometry, Result>', 'نتيجة تصميم لكل شكل من الأشكال الثلاثة'],
        ['comparison', 'GeometryComparisonReport', 'قرار المفاضلة والتوصية'],
        ['warnings', 'string[]', 'التحذيرات المتراكمة من كل المراحل'],
        ['status', 'enum', 'SUCCESS | PARTIAL | FAILED'],
        ['validationErrors', 'ValidationError[]', 'أخطاء التحقق إن وُجدت'],
    ],
    col_widths=[5, 4, 7], font_size=10)

add_page_break(doc)

print("✓ تم كتابة الفصول 1-3")

# ═══════════════════════════════════════════════════════════════════════
# الفصل الرابع: واجهات المصادقة وإدارة المستخدمين
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الرابع: واجهات المصادقة وإدارة المستخدمين', level=1, color=NAVY)

add_heading_ar(doc, '4.1 صفحة تسجيل الدخول /auth/login', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة تسجيل الدخول هي نقطة الدخول الرئيسية للمستخدمين المسجّلين. تتولى المصادقة '
    'الآمنة باستخدام bcrypt للتحقق من كلمة المرور، ثم إنشاء ملف تعريف جلسة مشفّر '
    'بـ AES-256-GCM يُخزَّن في cookie باسم sbv3-session. بعد التحقق، يُعاد توجيه '
    'المستخدم إلى لوحة التحكم أو إلى صفحة انتظار الموافقة إذا كانت حالته PENDING.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات (Inputs)', level=3)
add_table_simple(doc,
    ['الحقل', 'النوع', 'التحقق', 'الوصف'],
    [
        ['email', 'string', 'بريد إلكتروني صالح', 'البريد المسجّل للمستخدم'],
        ['password', 'string', '6 أحرف على الأقل', 'كلمة المرور (يُتحقق منها بـ bcrypt)'],
    ],
    col_widths=[3, 2, 4, 7], font_size=10)

add_heading_ar(doc, 'المخرجات (Outputs)', level=3)
add_para_ar(doc, 'بناءً على نتيجة التحقق، أحد السيناريوهات الثلاثة:')
add_bullet_ar(doc, 'نجاح + موافق عليه: ملف تعريف جلسة + إعادة توجيه إلى /dashboard')
add_bullet_ar(doc, 'نجاح + معلّق: إعادة توجيه إلى /auth/pending مع رسالة انتظار')
add_bullet_ar(doc, 'فشل: رسالة خطأ عربية واضحة (بدون كشف تفاصيل التقنية)')

add_heading_ar(doc, 'الحسابات والعمليات', level=3)
add_para_ar(doc, 'تتسلسل عمليات المصادقة كما يلي:')

steps_login = [
    ('1. التحقق من المدخلات', 'طبقة Zod تفحص صحة البريد الإلكتروني وطول كلمة المرور قبل أي معالجة.'),
    ('2. الاستعلام عن المستخدم', 'يتم استدعاء Supabase REST API: GET /rest/v1/users?email=eq.<email> لجلب سجل المستخدم.'),
    ('3. فحص حالة الاشتراك', 'إذا كان subscriptionStatus = REJECTED أو SUSPENDED يُرفض الدخول فوراً.'),
    ('4. التحقق من كلمة المرور', 'bcrypt.compare(password, user.passwordHash) — 12 rounds.'),
    ('5. إنشاء الجلسة', 'توليد sessionId عشوائي + تشفير AES-256-GCM للبيانات الحساسة باستخدام ENCRYPTION_KEY.'),
    ('6. تخزين الجلسة', 'يتم حفظها في قاعدة البيانات + إصدار cookie HttpOnly مع sameSite=lax.'),
    ('7. إعادة التوجيه', 'إلى /dashboard (للـ APPROVED) أو /auth/pending (للـ PENDING).'),
]
for title, desc in steps_login:
    add_bullet_ar(doc, f'{title}: {desc}')

add_formula_box(doc, 'bcrypt(password, salt) → hash  (cost factor = 12)', 'صيغة التحقق')
add_formula_box(doc, 'AES-256-GCM(sessionData, ENCRYPTION_KEY) → encrypted_payload', 'تشفير الجلسة')

add_info_box(doc, 'ملاحظة أمنية',
    ['كلمة المرور لا تُخزَّن مطلقاً بصيغة نصية. فقط hash بـ bcrypt يُحفظ في قاعدة البيانات. '
     'ملف تعريف الجلسة في cookie مشفّر بالكامل، ولا يحتوي على معلومات حساسة مكشوفة.'],
    color_bg='FCE4EC', color_border='C62828', icon='⚠')

add_page_break(doc)

add_heading_ar(doc, '4.2 صفحة التسجيل /auth/register', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة التسجيل تتيح لمستخدمين جدد إنشاء حساب في المنصة. لا تُمنح صلاحيات الوصول '
    'فوراً — كل حساب جديد يُنشأ بحالة اشتراك PENDING، ولا يُفعَّل إلا بعد موافقة '
    'المدير الحوكمي. هذا يضمن أن المنصة لا تُستخدم إلا من قبل جهات مخوّلة.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات (Inputs)', level=3)
add_table_simple(doc,
    ['الحقل', 'النوع', 'التحقق', 'الوصف'],
    [
        ['email', 'string', 'صيغة بريد + فريد', 'البريد الإلكتروني (لا يمكن تكراره)'],
        ['password', 'string', '8 أحرف + رقم + رمز', 'كلمة المرور (يُحفظ hash فقط)'],
        ['confirmPassword', 'string', 'تطابق password', 'تأكيد كلمة المرور'],
        ['fullName', 'string', '3 أحرف على الأقل', 'الاسم الكامل للمستخدم'],
        ['organization', 'string', 'اختياري', 'الجهة/المؤسسة'],
        ['role', 'enum', 'ENGINEER | REVIEWER | VIEWER', 'الدور المطلوب (يحتاج موافقة)'],
    ],
    col_widths=[3, 2, 4, 7], font_size=10)

add_heading_ar(doc, 'المخرجات (Outputs)', level=3)
add_bullet_ar(doc, 'حساب جديد بحالة subscriptionStatus = PENDING')
add_bullet_ar(doc, 'دور مؤقت: VIEWER (يُحدّث بعد الموافقة)')
add_bullet_ar(doc, 'أذونات فارغة (permissions = [])')
add_bullet_ar(doc, 'إعادة توجيه إلى /auth/pending مع رسالة انتظار')

add_heading_ar(doc, 'الحسابات والعمليات', level=3)
steps_register = [
    ('1. التحقق من المدخلات', 'Zod schema تشمل قواعد قوة كلمة المرور (8+ أحرف، رقم، رمز خاص).'),
    ('2. فحص التفرّد', 'GET /rest/v1/users?email=eq.<email> — إذا وُجد، يُرفض التسجيل.'),
    ('3. توليد الـ hash', 'bcrypt.hash(password, 12) → passwordHash'),
    ('4. إنشاء السجل', 'INSERT INTO users (email, passwordHash, fullName, role, subscriptionStatus=PENDING).'),
    ('5. تسجيل الحدث', 'INSERT INTO audit_logs (action=USER_REGISTERED, userId, timestamp).'),
    ('6. إعادة التوجيه', 'إلى /auth/pending.'),
]
for title, desc in steps_register:
    add_bullet_ar(doc, f'{title}: {desc}')

add_formula_box(doc, 'cost factor = 12  (bcrypt rounds)  →  ~250ms لكل تحقق', 'قوة bcrypt')

add_page_break(doc)

add_heading_ar(doc, '4.3 صفحة انتظار الموافقة /auth/pending', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة بسيطة تُعرض للمستخدمين الجدد بعد التسجيل، أو للمستخدمين الذين لم تتم '
    'موافقة المدير عليهم بعد. تعرض حالة الاشتراك الحالية ورسالة إيضاحية تشرح أن '
    'الوصول سيُفعَّل بعد المراجعة الحوكمية.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_para_ar(doc, 'تستقبل الصفحة معرّف الجلسة من cookie، وتستعلم عن حالة المستخدم.')

add_heading_ar(doc, 'المخرجات', level=3)
add_table_simple(doc,
    ['الحالة', 'الرسالة المعروضة', 'الإجراء المتاح'],
    [
        ['PENDING', 'طلبك قيد المراجعة من قبل المدير الحوكمي', 'انتظار فقط'],
        ['APPROVED', 'تمت الموافقة! يمكنك الدخول إلى لوحة التحكم', 'زر "تسجيل الدخول"'],
        ['REJECTED', 'تم رفض طلب الاشتراك. تواصل مع الإدارة', 'تواصل مع الإدارة'],
        ['SUSPENDED', 'تم إيقاف حسابك مؤقتاً', 'تواصل مع الإدارة'],
    ],
    col_widths=[3, 7, 6], font_size=10)

add_page_break(doc)

add_heading_ar(doc, '4.4 لوحة المدير /admin', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'لوحة المدير الحوكمي هي الأداة السيادية لإدارة المستخدمين والصلاحيات. '
    'يصل إليها فقط المستخدمون بدور ADMIN. تتضمن 4 تبويبات رئيسية: نظرة عامة، '
    'طلبات معلّقة، المستخدمون، وخط الأساس. يمكن للمدير من خلالها الموافقة على '
    'المستخدمين، تعليقهم، تغيير أدوارهم، إدارة 7 أذونات تفصيلية، وحذف الحسابات.',
    indent_first=True)

add_heading_ar(doc, 'الأذونات السبعة (RBAC)', level=3)
add_table_simple(doc,
    ['الإذن', 'الرمز', 'الوصف'],
    [
        ['1', 'engine:run', 'تشغيل محرك الحساب الكامل'],
        ['2', 'engine:read', 'قراءة نتائج المحرك'],
        ['3', 'reports:read', 'قراءة التقارير'],
        ['4', 'reports:export', 'تصدير التقارير PDF/Excel'],
        ['5', 'benchmark:run', 'تشغيل اختبارات BMK-01/02/03'],
        ['6', 'admin:manage_users', 'إدارة المستخدمين (للمدير المفوّض)'],
        ['7', 'admin:view_audit', 'عرض سجل التدقيق الكامل'],
    ],
    col_widths=[1.5, 5, 10], font_size=10)

add_heading_ar(doc, 'الإجراءات السيادية التسعة', level=3)
add_table_simple(doc,
    ['الإجراء', 'المدخلات', 'المخرجات', 'التأثير'],
    [
        ['seedDefaultAdmin', '—', 'حساب مدير افتراضي', 'تهيئة أولية'],
        ['getAllUsers', '—', 'مصفوفة users', 'قراءة فقط'],
        ['getPendingUsers', '—', 'users بحالة PENDING', 'قراءة فقط'],
        ['approveUser', 'userId', 'subscriptionStatus=APPROVED', 'تفعيل الوصول'],
        ['rejectUser', 'userId + reason', 'subscriptionStatus=REJECTED', 'رفض دائم'],
        ['suspendUser', 'userId + reason', 'subscriptionStatus=SUSPENDED', 'إيقاف مؤقت'],
        ['reactivateUser', 'userId', 'subscriptionStatus=APPROVED', 'إعادة تفعيل'],
        ['updateUserPermissions', 'userId + permissions[]', 'مصفوفة أذونات جديدة', 'تعديل RBAC'],
        ['changeUserRole', 'userId + newRole', 'role محدّث', 'تغيير الدور'],
        ['deleteUser', 'userId + confirm', 'حذف السجل', 'إزالة كاملة'],
    ],
    col_widths=[4, 4, 5, 4], font_size=9)

add_heading_ar(doc, 'سجل التدقيق (Audit Trail)', level=3)
add_para_ar(doc,
    'كل إجراء سيادي يُسجَّل في جدول audit_logs مع: timestamp, adminId, action, '
    'targetUserId, before_state, after_state, ipAddress, userAgent. هذا يضمن '
    'إمكانية تتبع كل تغيير في النظام، وهو متطلب أساسي للحوكمة السيادية.',
    indent_first=True)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════
# الفصل الخامس: محرك الاختراق
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الخامس: محرك الاختراق (Penetration Engine)', level=1, color=NAVY)

add_heading_ar(doc, '5.1 نظرة عامة', level=2)
add_para_ar(doc,
    'محرك الاختراق هو أول محرك حسابي في خط الأنابيب. يحسب عمق اختراق القنبلة '
    'في الوسط المستهدف (تربة أو خرسانة) باستخدام معادلات الأطروحة المرجعية '
    '(المعادلات 13-19). النتائج الرئيسية هي: عمق الاختراق x₁، عمق الحفرة h₀، '
    'الشحنة الفعّالة C_ef، والبعد الفعلي R_actual — وكلها قيم مقفلة تُستخدم '
    'في المحركات اللاحقة.',
    indent_first=True)

add_heading_ar(doc, '5.2 المدخلات (PenetrationInput)', level=2)
add_table_simple(doc,
    ['الحقل', 'النوع', 'الوحدة', 'النطاق', 'الوصف'],
    [
        ['weaponId', 'string', '—', '—', 'معرف السلاح من المكتبة (مثل MK83)'],
        ['impactVelocity', 'number', 'm/s', '50-1000', 'سرعة اصطدام القنبلة'],
        ['soilTypeCode', 'enum', '—', '6 قيم', 'SOFT_SOIL | MEDIUM_SOIL | HARD_ROCK | REINFORCED_SAND | CONCRETE | REINFORCED_CONCRETE'],
        ['impactAngleDeg', 'number?', 'درجة', '0-90', 'زاوية الاصطدام (0 = عمودي)'],
        ['detonationDelayMs', 'number?', 'ms', '≥ 0', 'تأخير التفجير (اختياري)'],
    ],
    col_widths=[3.5, 1.8, 1.8, 2.2, 7], font_size=9)

add_heading_ar(doc, '5.3 المخرجات (PenetrationOutput)', level=2)
add_table_simple(doc,
    ['الحقل', 'الوحدة', 'الوصف', 'الاستخدام اللاحق'],
    [
        ['penetrationDepth (x₁)', 'm', 'عمق الاختراق الكامل', 'حساب h_z, R_actual'],
        ['craterDepth (h₀)', 'm', 'عمق الحفرة المكافئ (= hz × 1.18)', 'تقييم الضرر'],
        ['requiredSpallingThickness', 'm', 'سماكة spalling = x₁ × 0.65', 'فحص السلامة'],
        ['lambda1', '—', 'معامل شكل الرأس (Eq.14)', 'محرك الاختراق'],
        ['lambda2', '—', 'معامل القطر (Eq.15)', 'محرك الاختراق'],
        ['nExp (n)', '—', 'أُس التأثير (Eq.16)', 'حساب τ'],
        ['cEffective (C_ef)', 'kg', 'الشحنة الفعّالة = 0.95 × K₁ × C', 'محرك الانفجار'],
        ['tsu (τ)', 'm', 'معامل زاوية الاختراق', 'حساب h_z'],
        ['hBarZ (h̄_z)', '—', 'العمق المختزل', 'حساب sigma_max'],
        ['warningMessages', 'string[]', 'تحذيرات ERR-PEN-* و WARN-PEN-*', 'واجهة المستخدم'],
    ],
    col_widths=[4, 1.5, 5, 5.5], font_size=9)

add_heading_ar(doc, '5.4 المعادلات الحسابية المرجعية', level=2)

add_heading_ar(doc, 'المعادلة (14): معامل شكل الرأس الحربي', level=3)
add_formula_box(doc, 'λ₁ = 0.5 + 0.4 × (Lh/D)^(2/3)')
add_para_ar(doc,
    'حيث Lh/D هو نسبة طول الرأس الحربي إلى القطر. هذه المعادلة تأخذ بالاعتبار '
    'تأثير شكل الرأس على قدرة الاختراق — الرؤوس الحادة نسبياً (نسبة عالية) '
    'تعطي قيمة λ₁ أعلى وبالتالي اختراق أعمق.')

add_heading_ar(doc, 'المعادلة (15): معامل تأثير القطر', level=3)
add_formula_box(doc, 'λ₂ = 2.8 × d^(1/3) − 1.3 × d^(1/2)')
add_para_ar(doc,
    'حيث d قطر القنبلة بالأمتار. هذه المعادلة تمثّل تأثير القطر على الاختراق — '
    'القنابل ذات الأقطار الكبيرة تختبر أقل نسبياً بسبب زيادة المساحة المعرضة.')

add_heading_ar(doc, 'المعادلة (16): أُس التأثير', level=3)
add_formula_box(doc, 'n = 3.5 − (Lh/D)')

add_heading_ar(doc, 'المعادلة (19): الشحنة الفعّالة', level=3)
add_formula_box(doc, 'C_ef = 0.95 × K₁ × C')
add_para_ar(doc,
    'حيث K₁ هو معامل التحويل إلى TNT المكافئ (يختلف حسب نوع المتفجر)، '
    'و C هو وزن الشحنة المتفجرة بالكيلوغرام. العامل 0.95 يمثّل كفاءة التفجير.')

add_heading_ar(doc, 'المعادلة (17): معامل زاوية الاختراق — قنبلة خارقة', level=3)
add_formula_box(doc, 'τ = 0.5 × lₖ × cos((α + n×α)/2)')

add_heading_ar(doc, 'المعادلة (18): معامل زاوية الاختراق — قنبلة انفجارية', level=3)
add_formula_box(doc, 'τ = 0.5 × dₖ')

add_heading_ar(doc, 'المعادلة (13): عمق الاختراق في التربة', level=3)
add_formula_box(doc, 'x₁ = λ₁ × λ₂ × Kpr × (P/d²) × V × cos(α)')
add_para_ar(doc,
    'حيث: Kpr = معامل اختراق التربة (من جدول SOIL_TABLE)، P = وزن القنبلة (kg)، '
    'd = القطر (m)، V = السرعة (m/s)، α = زاوية الاصطدام. هذه المعادلة الجوهرية '
    'تجمع بين خصائص السلاح (P, d, V) وخصائص الوسط (Kpr) لتقدير عمق الاختراق.')

add_heading_ar(doc, '5.5 الحسابات المشتقّة', level=2)

add_heading_ar(doc, 'العمق الصافي والعمق المختزل', level=3)
add_formula_box(doc, 'h_z = max(x₁ − τ, 0)')
add_formula_box(doc, 'craterDepth (h₀) = h_z × 1.18')
add_formula_box(doc, 'h̄_z = h_z / ∛(C_ef)')
add_para_ar(doc,
    'العمق الصافي h_z يخصم من x₁ قيمة τ (طول الجزء غير المخترق). عمق الحفرة h₀ '
    'يأخذ معامل 1.18 المرجعي. العمق المختزل h̄_z يُستخدم في حساب σ_max في محرك الضغط.')

add_heading_ar(doc, 'سماكة Spalling المطلوبة', level=3)
add_formula_box(doc, 'requiredSpallingThickness = x₁ × 0.65')

add_heading_ar(doc, '5.6 التحذيرات التلقائية', level=2)
add_table_simple(doc,
    ['الكود', 'الشرط', 'الإجراء'],
    [
        ['ERR-PEN-01', 'x₁ ≤ 0', 'خطأ: تحقق من المدخلات'],
        ['WARN-PEN-01', 'h_z ≤ 0', 'تحذير: القنبلة لم تخترق السقف'],
        ['WARN-PEN-02', 'h̄_z < 0.7', 'تحذير: تخفيض σ_max بعامل 0.6'],
        ['WARN-PEN-03', 'V > 400 m/s', 'تحذير: نطاق سرعة عالية — تحقق من المعادلة'],
    ],
    col_widths=[3, 5, 8], font_size=10)

add_heading_ar(doc, '5.7 القيم المرجعية BMK-02', level=2)
add_para_ar(doc,
    'في الحالة المرجعية BMK-02 (قنبلة MK83 في تربة MEDIUM_SOIL، سرعة 350 m/s، '
    'زاوية 20°)، تنتج المنصة القيم التالية التي تُستخدم كنقطة تحقق:')

add_table_simple(doc,
    ['الرمز', 'القيمة المرجعية', 'الوحدة', 'الوصف'],
    [
        ['C_ef', '334.766', 'kg', 'الشحنة الفعّالة'],
        ['λ₁', '1.1347', '—', 'معامل شكل الرأس'],
        ['λ₂', '1.2125', '—', 'معامل القطر'],
        ['h_pr (x₁)', '3.6495', 'm', 'عمق الاختراق'],
        ['R_actual', '7.6231', 'm', 'البعد الفعلي'],
        ['h_z', '2.6895', 'm', 'العمق الصافي'],
        ['h̄_z', '0.3881', '—', 'العمق المختزل'],
        ['Zp', '5.8213', 'm/kg^1/3', 'البعد المختزل'],
    ],
    col_widths=[3, 4, 3, 6], font_size=10)

add_page_break(doc)

print("✓ تم كتابة الفصل 4-5")

# ═══════════════════════════════════════════════════════════════════════
# الفصل السادس: محرك الضغط العصفي
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل السادس: محرك الضغط العصفي (Blast Pressure Engine)', level=1, color=NAVY)

add_heading_ar(doc, '6.1 نظرة عامة', level=2)
add_para_ar(doc,
    'محرك الضغط العصفي يحسب الضغط الناتج عن الانفجار على المنشأ، باستخدام صيغ '
    'سادوفسكي الشهيرة للانفجارات في الهواء، مع تعديلات للتربة (صيغة الإجهاد '
    'σ_max = A × Z^(-n1) من جداول الاستيفاء I-3). يُنتج المحرك الضغط الجانبي '
    'Pso، الضغط المنعكس Pr، الضغط التصميمي P_design، وزمن الطور الموجب τ⁺ '
    'والزمن الفعال τ_eff.',
    indent_first=True)

add_heading_ar(doc, '6.2 النسختان من المحرك', level=2)
add_para_ar(doc,
    'يوجد نسختان متكاملتان من محرك الضغط: (1) blast-pressure-core.ts للحساب '
    'المبسّط السريع في Orchestrator العام، و(2) blast-loads.ts للحساب التفصيلي '
    'في صفحة /dashboard/blast-loads مع مسارين منفصلين (سقف + جدار) وتحقق كامل '
    'من القيم المرجعية BMK-02. النسخة الثانية تستخدم القيم المقفلة من الخطوات 2-4 '
    'وتعالج الانحرافات عن المرجع آلياً.',
    indent_first=True)

add_heading_ar(doc, '6.3 المدخلات (BlastInput)', level=2)
add_table_simple(doc,
    ['الحقل', 'النوع', 'الوحدة', 'النطاق', 'الوصف'],
    [
        ['equivalentTNTWeight', 'number', 'kg', '0-50000', 'وزن TNT المكافئ (C_ef من محرك الاختراق)'],
        ['radialDistance', 'number', 'm', '0-1000', 'البعد الشعاعي عن مركز الانفجار'],
        ['soilTypeCode', 'enum?', '—', '6 قيم', 'نوع التربة (لجداول I-3)'],
        ['ceilingDepth', 'number?', 'm', '> 0', 'عمق السقف فوق مركز الانفجار'],
    ],
    col_widths=[4, 1.8, 1.8, 2.2, 6.5], font_size=9)

add_heading_ar(doc, '6.4 المخرجات (BlastOutput)', level=2)
add_table_simple(doc,
    ['الحقل', 'الوحدة', 'الوصف'],
    [
        ['scaledDistanceZ (Z)', 'm/kg^1/3', 'البعد المختزل = R / ∛C'],
        ['pSideOnMpa (Pso)', 'MPa', 'الضغط الجانبي (Sadovsky)'],
        ['pReflectedMpa (Pr)', 'MPa', 'الضغط المنعكس = 2 × Kp × σ_max'],
        ['durationMs', 'ms', 'مدة الطور الموجب × 1000'],
        ['pDesignMpa', 'MPa', 'الضغط التصميمي = P_static + P_equivalent'],
        ['pDesignKPa', 'kPa', 'الضغط التصميمي × 1000 (للتصميم الإنشائي)'],
        ['sigmaMaxMpa (σ_max)', 'MPa', 'الإجهاد الأقصى في التربة = A × Z^(-n1)'],
        ['tauPlus (τ⁺)', 's', 'زمن الطور الموجب (Sadovsky)'],
        ['tauEffective (τ_eff)', 's', 'الزمن الفعال = τ⁺ × f(ΔPmax)'],
        ['rCritical (R_critical)', 'm', 'البعد الحرج = 1.1 × ∛C'],
        ['dynamicConditionMet', 'boolean', 'تحقق شرط الديناميكية: τ_eff ≥ 0.2π/ω'],
        ['coreConditionMet', 'boolean', 'تحقق شرط نواة المقطع'],
    ],
    col_widths=[5, 2, 9], font_size=10)

add_heading_ar(doc, '6.5 المعادلات الحسابية المرجعية', level=2)

add_heading_ar(doc, 'المعادلة (1): الضغط الزائد الساقط — سادوفسكي', level=3)
add_formula_box(doc, 'ΔP = 0.1 × ∛C / R + 0.43 × ∛(C²) / R² + 1.4 × C / R³  [MPa]')
add_para_ar(doc,
    'حيث C بالكيلوغرام (وزن TNT المكافئ) و R بالأمتار (البعد الشعاعي). '
    'الناتج بالميغاباسكال. هذه الصيغة تُعطي الضغط الجانبي Pso للموجة الانفجارية '
    'في الهواء الحر.')

add_heading_ar(doc, 'المعادلة (4): زمن الطور الموجب', level=3)
add_formula_box(doc, 'τ⁺ = 1.7 × 10⁻³ × ∛(√C) × √R  [s]')

add_heading_ar(doc, 'المعادلة (5): الاندثار (Impulse)', level=3)
add_formula_box(doc, 'I = 6.3 × ∛(C²) / R')

add_heading_ar(doc, 'المعادلة (8): الزمن الفعال — دالة الاستيفاء I-9', level=3)
add_formula_box(doc, 'τ_eff = τ⁺ × f(ΔPmax)')
add_formula_box(doc, 'f(ΔPmax) = 0.0008 × ΔPmax² − 0.0384 × ΔPmax + 1.0013')

add_heading_ar(doc, 'الإجهاد الأقصى في التربة', level=3)
add_formula_box(doc, 'σ_max = A × Z^(-n1)  [MPa]')
add_para_ar(doc,
    'حيث Z هو البعد المختزل R / ∛C، و A و n1 هما معاملات الاستيفاء من الجدول I-3 '
    'التي تعتمد على نوع التربة. مثلاً للطين المتوسط A=18, n1=2.8؛ '
    'للطين مع حجارة A=5, n1=3.')

add_heading_ar(doc, 'الضغط المنعكس', level=3)
add_formula_box(doc, 'Pr = 2 × Kp × σ_max  (Kp = 0.86 مرجعي)')

add_heading_ar(doc, 'البعد الحرج', level=3)
add_formula_box(doc, 'R_critical = 1.1 × ∛C  (مرجعي للتربة الطينية)')

add_heading_ar(doc, 'شرط الديناميكية', level=3)
add_formula_box(doc, 'dynamicConditionMet = (τ_eff ≥ 0.2 × π / ω)')
add_para_ar(doc,
    'حيث ω هو التردد الدائري الطبيعي للوحة (يُحسب من خصائصها الديناميكية). '
    'إذا تحقق الشرط، معامل الديناميكية Kd = 0.9، وإلا Kd = 0 (حمل ساكن).')

add_heading_ar(doc, 'الضغط التصميمي النهائي', level=3)
add_formula_box(doc, 'P_design = P_static + P_equivalent')
add_formula_box(doc, 'P_equivalent = Kd × P_max  (حيث P_max = Pr)')
add_formula_box(doc, 'P_static = ceilingDepth × 2000 / 10000 × 0.0980665  [MPa]')

add_heading_ar(doc, '6.6 المعادلات الإضافية (blast-loads.ts)', level=2)
add_para_ar(doc,
    'النسخة التفصيلية blast-loads.ts تضيف معادلات إضافية للمسارين المنفصلين:')

add_formula_box(doc, 'h̄ = Hp_cm / (100 × ∛C_ef)', 'النسبة المختزلة للسماكة')
add_formula_box(doc, 'R_ekv = R_actual − h_pr × (1 − Hp_cm/ht_cm)', 'البعد المكافئ')
add_formula_box(doc, 'R* = R̄ × ∛C_ef', 'البعد النجمي')
add_formula_box(doc, 'τ = 1.7×10⁻³ × ∛(√C_ef) × √R_ekv', 'زمن الطور الموجب (سادوفسكي)')
add_formula_box(doc, 'ω = (π²/L²) × √(Ea × I / (mass_density × A))', 'التردد الدائري')
add_formula_box(doc, 'I = 100 × Hp³ / 12  [cm⁴]', 'عزم القصور الذاتي')
add_formula_box(doc, 'P_max = Kpod × σ_max × Kp', 'الضغط الأقصى (من الإجهاد)')
add_formula_box(doc, 'P_ekv = Kd × kpsi × P_max', 'الضغط المكافئ')
add_formula_box(doc, 'Pp = P_ekv + Pct', 'ضغط التصميم النهائي')

add_para_ar(doc,
    'حيث kpsi = 0.9 للسقف و 0.85 للجدار (معامل المسار)، و Pct هو الضغط الثابت '
    'من وزن التربة فوق المنشأ. للسقف: Pct = γ_g × (Z − Hp) / 10000. '
    'للجدار: Pct = γ_g × H × Ka / 10000 (حيث Ka ≈ 0.5).')

add_heading_ar(doc, '6.7 القيم المرجعية BMK-02', level=2)
add_para_ar(doc, 'للحالة المرجعية، النتائج في كلا المسارين:')

add_heading_ar(doc, 'مسار السقف (STEP5_ROOF)', level=3)
add_table_simple(doc,
    ['الرمز', 'القيمة', 'الوحدة', 'الوصف'],
    [
        ['h̄', '0.1180', '—', 'النسبة المختزلة للسماكة'],
        ['R_ekv', '6.1162', 'm', 'البعد المكافئ'],
        ['R*', '2.4255', 'm', 'البعد النجمي'],
        ['max_bv', '2.8803', 'm/s', 'أقصى سرعة جسيمية'],
        ['τ', '0.2650', 's', 'زمن الطور الموجب'],
        ['τ_ef', '0.2378', 's', 'الزمن الفعال'],
        ['τ_n', '0.0365', 's', 'الزمن الطبيعي'],
        ['ω', '561.67', 'rad/s', 'التردد الدائري'],
        ['C_dyn', '46.81', 'm/s', 'السرعة الديناميكية'],
        ['μ_struct', '0.8862', '—', 'نسبة المطاوعة الإنشائية'],
        ['P_max', '4.6084', 'kg/cm²', 'الضغط الأقصى'],
        ['P_ekv', '3.8158', 'kg/cm²', 'الضغط المكافئ'],
        ['Pct', '1.1053', 'kg/cm²', 'الضغط الثابت'],
        ['Pp', '4.9211', 'kg/cm²', 'ضغط التصميم النهائي'],
    ],
    col_widths=[3, 3.5, 3, 6.5], font_size=9)

add_heading_ar(doc, 'مسار الجدار (STEP5_WALL)', level=3)
add_table_simple(doc,
    ['الرمز', 'القيمة', 'الوحدة', 'الوصف'],
    [
        ['Pp_wall', '3.7845', 'kg/cm²', 'ضغط التصميم للجدار'],
        ['kpsi', '0.85', '—', 'معامل المسار (أقل من السقف)'],
    ],
    col_widths=[3, 3.5, 3, 6.5], font_size=9)

add_page_break(doc)

print("✓ تم كتابة الفصل 6")

# ═══════════════════════════════════════════════════════════════════════
# الفصل السابع: محرك التصميم الإنشائي
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل السابع: محرك التصميم الإنشائي (Structural Concrete Engine)', level=1, color=NAVY)

add_heading_ar(doc, '7.1 نظرة عامة', level=2)
add_para_ar(doc,
    'محرك التصميم الإنشائي يحسب السماكة المطلوبة للمقطع الخرساني ومساحة التسليح، '
    'ثم يتحقق من سلامة المقطع وفق الكود السوري 2024 ودليل UFC 3-340-02. '
    'يستقبل الضغط التصميمي P_design من محرك الضغط العصفي، ويُنتج تقرير تحقق '
    'شامل يشمل: السماكة، التسليح، نسبة المطاوعة، فحص اللامركزية، وفحص القص الثاقب.',
    indent_first=True)

add_heading_ar(doc, '7.2 المدخلات (DesignInput)', level=2)
add_table_simple(doc,
    ['الحقل', 'النوع', 'الوحدة', 'النطاق', 'الوصف'],
    [
        ['pDesignMpa', 'number', 'MPa', '0-500', 'الضغط التصميمي (من محرك الانفجار)'],
        ['geometryType', 'enum', '—', '3 قيم', 'RECTANGULAR | CIRCULAR | ARCHED'],
        ['tunnelSpanShort (ap)', 'number', 'm', '0-30', 'البحر القصير للنفق'],
        ['tunnelSpanLong (bp)', 'number', 'm', '0-50', 'البحر الطويل للنفق'],
        ['fcMpa', 'number', 'MPa', '15-80', 'مقاومة ضغط الخرسانة'],
        ['fyMpa', 'number', 'MPa', '200-600', 'إجهاد خضوع الحديد'],
        ['slabThicknessHintMm', 'number?', 'mm', '> 0', 'سماكة مقترحة (اختياري)'],
        ['reinforcementRatio', 'number?', '—', '0.001-0.04', 'نسبة تسليح مقترحة'],
    ],
    col_widths=[3.5, 1.8, 1.8, 2.2, 6.5], font_size=9)

add_heading_ar(doc, '7.3 المخرجات (SectionDesignResult)', level=2)
add_table_simple(doc,
    ['الحقل', 'الوحدة', 'الوصف'],
    [
        ['requiredThicknessMeters', 'm', 'السماكة المطلوبة للمقطع'],
        ['requiredSteelAreaCm2PerMeter', 'cm²/m', 'مساحة التسليح المطلوبة للمتر الطولي'],
        ['ductilityRatio', '—', 'نسبة المطاوعة = fy / fc'],
        ['validation.status', 'enum', 'SUCCESS | WARNING | FAILURE'],
        ['validation.eccentricityRatio', '—', 'نسبة اللامركزية e/h'],
        ['validation.punchingShearRatio', '—', 'إجهاد القص الفعلي/المسموح'],
        ['validation.reinforcementRatio', '—', 'نسبة التسليح الفعلية'],
        ['validation.failures', 'string[]', 'قائمة الإخفاقات التفصيلية'],
        ['validation.explanation', 'string', 'شرح قرار التحقق'],
        ['validation.ruleId', 'string', 'معرف القاعدة الحاكمة (SYR-2024-OK / REVIEW)'],
    ],
    col_widths=[5, 2, 9], font_size=10)

add_heading_ar(doc, '7.4 المعادلات الحسابية', level=2)

add_heading_ar(doc, 'تطبيق معاملات التضخيم الديناميكي (DIF)', level=3)
add_formula_box(doc, 'fcd = fc × DIF_concrete_compression  (1.25 وفق UFC)')
add_formula_box(doc, 'fsd = fy × DIF_steel_tension  (1.20 وفق UFC)')
add_para_ar(doc,
    'معاملات التضخيم الديناميكي (Dynamic Increase Factors) تأخذ بالاعتبار زيادة '
    'مقاومة المواد تحت الأحمال السريعة. للخرسانة في الضغط: 1.25، للحديد في الشد: 1.20، '
    'للقص الثاقب: 1.25.')

add_heading_ar(doc, 'العزم الديناميكي', level=3)
add_formula_box(doc, 'P_design_kPa = P_design_MPa × 1000')
add_formula_box(doc, 'M = P_design_kPa × ap² / 8  [kN.m]  (لوح ببساطة بسيطة)')
add_para_ar(doc,
    'يُحسب العزم الديناميكي بفرض لوح بسيط ببحرين (بسيط من طرفين) تحت حمل موزّع '
    'منتظم يساوي الضغط التصميمي. المعامل 1/8 هو معامل اللوح المرن البسيط.')

add_heading_ar(doc, 'العمق الفعلي', level=3)
add_formula_box(doc, 'd_eff = h − cover  (cover = 50mm وفق الكود السوري)')

add_heading_ar(doc, 'السماكة المطلوبة', level=3)
add_formula_box(doc, 'Rb = fc × 0.9 × 10  [kg/cm²]  (تحويل تقريبي)')
add_formula_box(doc, 'h₀ = √(M × 10⁵ / (Rb × b × 0.3))')
add_formula_box(doc, 'h_required = h₀ + cover  [mm]')

add_heading_ar(doc, 'مساحة التسليح المطلوبة', level=3)
add_formula_box(doc, 'As = (M × 10⁶) / (fsd × 0.875 × d_eff)  [mm²]')
add_para_ar(doc,
    'حيث 0.875 هو معامل ذراع العزم التقريبي (z = 0.875 × d_eff). الناتج يُحوَّل '
    'إلى cm²/m بقسمة على 100.')

add_heading_ar(doc, 'نسبة التسليح', level=3)
add_formula_box(doc, 'ρ = max(As / (d_eff × 100), ρ_min)  (ρ_min = 0.0025)')

add_heading_ar(doc, 'نسبة المطاوعة', level=3)
add_formula_box(doc, 'ductilityRatio = fy / fc')

add_heading_ar(doc, '7.5 فحوصات التحقق', level=2)

add_heading_ar(doc, 'فحص اللامركزية (Eccentricity)', level=3)
add_formula_box(doc, 'N = P_design_kPa × ap × bp  (القوة المحورية)')
add_formula_box(doc, 'e = M × 10⁶ / (N × 10³)  [mm]')
add_formula_box(doc, 'e_limit = h / 6  (حد النواة)')
add_formula_box(doc, 'ERR-CORE-01:  e > h/6  →  فشل')

add_heading_ar(doc, 'فحص القص الثاقب (Punching Shear)', level=3)
add_formula_box(doc, 'b₀ = 2 × (b_col + d_eff) + 2 × (h_col + d_eff)  (المحيط الحرج)')
add_formula_box(doc, 'A_critical = (b_col + d_eff) × (h_col + d_eff) / 10⁶  [m²]')
add_formula_box(doc, 'V_actual = P_design_kPa × (A_tributary − A_critical)  [kN]')
add_formula_box(doc, 'v_actual = V_actual × 1000 / (b₀ × d_eff)  [MPa]')
add_formula_box(doc, 'v_cd = 0.25 × √(fcd)  [MPa]  (الكود السوري)')
add_formula_box(doc, 'ERR-PUNCH-01:  v_actual > v_cd  →  فشل')

add_heading_ar(doc, 'فحوصات إضافية', level=3)
add_bullet_ar(doc, 'ERR-MAT-01: fc < 25 MPa — مقاومة خرسانة أقل من الحد الأدنى')
add_bullet_ar(doc, 'ERR-RHO-01: ρ < ρ_min — نسبة تسليح أقل من الحد الأدنى')

add_heading_ar(doc, '7.6 قرار التحقق النهائي', level=3)
add_table_simple(doc,
    ['الحالة', 'الشرط', 'النتيجة'],
    [
        ['SUCCESS', 'لا فشل', 'المقطع يحقق متطلبات الكود'],
        ['WARNING', 'فشل في فحوصات غير حرجة', 'المقطع يحتاج مراجعة'],
        ['FAILURE', 'فشل في ERR-PUNCH أو ERR-CORE', 'المقطع غير آمن — إعادة تصميم'],
    ],
    col_widths=[3, 6, 7], font_size=10)

add_heading_ar(doc, '7.7 ثوابت الكود', level=2)
add_table_simple(doc,
    ['الثابت', 'القيمة', 'المصدر'],
    [
        ['DIF_CONCRETE_COMPRESSION', '1.25', 'UFC 3-340-02'],
        ['DIF_CONCRETE_TENSION', '1.25', 'UFC 3-340-02'],
        ['DIF_STEEL_TENSION', '1.20', 'UFC 3-340-02'],
        ['DIF_PUNCHING_SHEAR', '1.25', 'UFC 3-340-02'],
        ['V_CD_COEFF', '0.25', 'الكود السوري 2024'],
        ['PHI_V', '0.85', 'الكود السوري 2024'],
        ['RHO_MIN', '0.0025', 'الكود السوري 2024'],
        ['RHO_MAX_COEFF', '0.5', 'الكود السوري 2024'],
        ['SAFETY_FACTOR_KN', '1.4', 'الكود السوري 2024'],
        ['COVER_MIN_MM', '50', 'الكود السوري 2024'],
    ],
    col_widths=[6, 3, 7], font_size=10)

add_page_break(doc)

print("✓ تم كتابة الفصل 7")

# ═══════════════════════════════════════════════════════════════════════
# الفصل الثامن: محرك تصميم التسليح
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الثامن: محرك تصميم التسليح (Rebar Design Engine)', level=1, color=NAVY)

add_heading_ar(doc, '8.1 نظرة عامة', level=2)
add_para_ar(doc,
    'محرك تصميم التسليح يأخذ مخرجات محرك التصميم الإنشائي (العزم البلاستيكي Mp '
    'والعمق الفعلي h₀) ويُفصّل التسليح: عدد القضبان، القطر الأمثل، مساحة التسليح '
    'المقدمة As_provided، والفحوصات الكاملة وفق الكود السوري 2024 و UFC 3-340-02 '
    'للأحمال الديناميكية. يعمل المحرك لكلا المسارين (سقف + جدار) بشكل مستقل.',
    indent_first=True)

add_heading_ar(doc, '8.2 المدخلات (RebarDesignInput)', level=2)
add_table_simple(doc,
    ['الحقل', 'النوع', 'الوحدة', 'الوصف'],
    [
        ['Mp', 'number', 'kg.cm', 'العزم البلاستيكي'],
        ['h0', 'number', 'cm', 'العمق الفعلي'],
        ['b', 'number', 'cm', 'عرض المقطع (عادة 100 cm للمتر الطولي)'],
        ['Rsd', 'number', 'kg/cm²', 'مقاومة القص الديناميكية للحديد'],
        ['Rbd', 'number', 'kg/cm²', 'مقاومة الانحناء الديناميكية للخرسانة'],
        ['h', 'number', 'cm', 'السماكة الكلية للمقطع'],
        ['path', 'enum', '—', '"roof" للسقف | "wall" للجدار'],
        ['mu_struct', 'number', '—', 'نسبة المطاوعة الإنشائية'],
    ],
    col_widths=[3, 2, 2, 9], font_size=10)

add_heading_ar(doc, '8.3 المخرجات (RebarDesignOutput)', level=2)
add_table_simple(doc,
    ['الحقل', 'الوحدة', 'الوصف'],
    [
        ['As_required', 'cm²/m', 'مساحة التسليح المطلوبة (محسوبة)'],
        ['As_provided', 'cm²/m', 'مساحة التسليح المقدمة (مختارة)'],
        ['rho (ρ)', '—', 'نسبة التسليح الفعلية = As / (b × h0)'],
        ['rho_min', '—', 'الحد الأدنى لنسبة التسليح'],
        ['rho_max', '—', 'الحد الأقصى لنسبة التسليح'],
        ['xi (ξ)', '—', 'العمق النسبي لمنطقة الضغط'],
        ['xi_max', '—', 'العمق النسبي الأقصى المسموح'],
        ['alpha_m', '—', 'معامل العزم النسبي = ξ × (1 - 0.5ξ)'],
        ['barCount', '—', 'عدد قضبان التسليح'],
        ['barDiameter', 'mm', 'قطر القضيب المختار'],
        ['singleBarArea', 'cm²', 'مساحة قضيب واحد'],
        ['minReinforcementOk', 'boolean', 'تحقق ρ ≥ ρ_min'],
        ['maxReinforcementOk', 'boolean', 'تحقق ρ ≤ ρ_max'],
        ['xiConditionOk', 'boolean', 'تحقق ξ ≤ ξ_max'],
        ['status', 'enum', 'OK | WARNING | FAILURE'],
        ['notes', 'string[]', 'ملاحظات وتحذيرات'],
    ],
    col_widths=[5, 2, 9], font_size=10)

add_heading_ar(doc, '8.4 المعادلات الحسابية', level=2)

add_heading_ar(doc, 'مساحة التسليح المطلوبة — حل تكراري', level=3)
add_formula_box(doc, 'As = Mp / (Rsd × z)  حيث  z = h0 − 0.5 × x')
add_formula_box(doc, 'x = Rsd × As / (Rbd × b)  (عمق منطقة الضغط)')
add_para_ar(doc,
    'الحل التكراري: نفترض z = 0.875 × h0 كتقريب أولي، نحسب As، نحسب x، نحسب z الجديد، '
    'نعيد الحساب حتى التقارب (تسامح 0.001 أو 10 تكرارات كحد أقصى).')

add_heading_ar(doc, 'العمق النسبي لمنطقة الضغط', level=3)
add_formula_box(doc, 'ξ = Rsd × As / (Rbd × b × h0)')

add_heading_ar(doc, 'معامل العزم النسبي', level=3)
add_formula_box(doc, 'αm = ξ × (1 − 0.5 × ξ)')

add_heading_ar(doc, 'نسبة التسليح', level=3)
add_formula_box(doc, 'ρ = As / (b × h0)')

add_heading_ar(doc, 'حدود التسليح', level=3)
add_formula_box(doc, 'ρ_min = max(0.0025, 0.26 × fctm / fyk)  للظروف العادية')
add_formula_box(doc, 'ρ_min = 0.003  (0.3%)  للأحمال الديناميكية')
add_formula_box(doc, 'ρ_max = 0.025  (2.5%)  للظروف العادية')
add_formula_box(doc, 'ρ_max = 0.04  (4%)  للأحمال الديناميكية (UFC 3-340-02)')

add_heading_ar(doc, 'العمق النسبي الأقصى المسموح', level=3)
add_formula_box(doc, 'ξ_max = 0.55  للكود السوري (ظروف عادية)')
add_formula_box(doc, 'ξ_max = 0.65  للأحمال الديناميكية (UFC 3-340-02)')

add_heading_ar(doc, '8.5 اختيار القضبان الأمثل', level=2)
add_para_ar(doc,
    'يختار المحرك تلقائياً القضبان الأمثل من جدول الأقطار القياسية (10، 12، 14، '
    '16، 18، 20، 22، 25، 28، 32 ملم) بحيث يحقق المساحة المطلوبة As_required مع '
    'أقل فائض ممكن، ومع مراعاة شرط التباعد الأدنى (5 cm بين القضبان).',
    indent_first=True)

add_table_simple(doc,
    ['القطر (mm)', 'المساحة (cm²)', 'القطر (mm)', 'المساحة (cm²)'],
    [
        ['10', '0.785', '22', '3.801'],
        ['12', '1.131', '25', '4.909'],
        ['14', '1.539', '28', '6.158'],
        ['16', '2.011', '32', '8.042'],
        ['18', '2.545', '—', '—'],
        ['20', '3.142', '—', '—'],
    ],
    col_widths=[3, 3, 3, 3], font_size=10)

add_heading_ar(doc, '8.6 الفحوصات والقرار', level=2)
add_table_simple(doc,
    ['الفحص', 'الشرط', 'الحالة'],
    [
        ['minReinforcementOk', 'ρ ≥ ρ_min', 'إذا فشل → WARNING'],
        ['maxReinforcementOk', 'ρ ≤ ρ_max', 'إذا فشل → FAILURE'],
        ['xiConditionOk', 'ξ ≤ ξ_max', 'إذا فشل → FAILURE'],
        ['status: OK', 'كل الفحوصات ناجحة', 'يحقق متطلبات الكود'],
        ['status: WARNING', 'فقط minReinforcement فشل', 'يستخدم ρ_min'],
        ['status: FAILURE', 'maxReinforcement أو ξ فشل', 'المقطع يحتاج توسيع'],
    ],
    col_widths=[4, 5, 7], font_size=10)

add_heading_ar(doc, '8.7 المعادلات الإضافية للحسابات الديناميكية', level=2)
add_formula_box(doc, 'Rsd = RsH × DIF_steel × n0  (DIF=1.2, n0=1.25)')
add_formula_box(doc, 'Rbd = RbH × DIF_concrete × n0 / 10  (DIF=1.25, n0=1.25)')
add_formula_box(doc, 'μ_struct = (RsH / RbH) × αm = (RsH / RbH) × ξ × (1 − 0.5ξ)')

add_heading_ar(doc, '8.8 القيم المرجعية', level=2)
add_para_ar(doc, 'في الحالة المرجعية BMK-02:')

add_table_simple(doc,
    ['الرمز', 'القيمة', 'الوحدة', 'الوصف'],
    [
        ['Rsd', '3937.5', 'kg/cm²', 'مقاومة القص الديناميكية للحديد'],
        ['Rbd', '236', 'kg/cm²', 'مقاومة الانحناء الديناميكية للخرسانة'],
        ['μ_struct (سقف)', '0.8862', '—', 'نسبة المطاوعة الإنشائية للسقف'],
        ['eta', '1.25', '—', 'معامل الكفاءة'],
    ],
    col_widths=[3, 3.5, 3, 6.5], font_size=10)

add_page_break(doc)

print("✓ تم كتابة الفصل 8")

# ═══════════════════════════════════════════════════════════════════════
# الفصل التاسع: محرك المفاضلة الهندسية
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل التاسع: محرك المفاضلة الهندسية (Geometry Comparison Engine)', level=1, color=NAVY)

add_heading_ar(doc, '9.1 نظرة عامة', level=2)
add_para_ar(doc,
    'محرك المفاضلة الهندسية هو المحرك الختامي في خط الأنابيب. يستقبل نتائج التصميم '
    'الإنشائي للأشكال الثلاثة (مستطيل، دائري، مقوّس) ويُقرّر الشكل الإنشائي الأمثل '
    'بناءً على 4 معايير مرجّحة: السماكة، وزن الحديد، الأمان، ونسبة المطاوعة. '
    'يُنتج تقرير مقارنة شامل مع مصفوفة نتائج وتوصية نهائية.',
    indent_first=True)

add_heading_ar(doc, '9.2 الأشكال الهندسية الثلاثة', level=2)
add_table_simple(doc,
    ['الشكل', 'الرمز', 'المميزات', 'الاستخدام النموذجي'],
    [
        ['مستطيل', 'RECTANGULAR', 'سهل التنفيذ + اقتصادي', 'أنفاق ببحر صغير/متوسط'],
        ['دائري', 'CIRCULAR', 'أفضل توزيع إجهاد + أقل سماكة', 'أنفاق عميقة + ضغط عالٍ'],
        ['مقوّس', 'ARCHED', 'توازن بين الكلفة والأداء', 'أنفاق ببحر كبير + أحمال متوسطة'],
    ],
    col_widths=[2.5, 3, 5, 5.5], font_size=10)

add_heading_ar(doc, '9.3 المدخلات والمخرجات', level=2)
add_para_ar(doc, 'يستقبل المحرك كائن SectionDesignResult لكل شكل من الأشكال الثلاثة، '
                  'ويُعيد GeometryComparisonReport يحتوي على:')

add_table_simple(doc,
    ['الحقل', 'النوع', 'الوصف'],
    [
        ['recommendedGeometry', 'enum', 'الشكل الموصى به (RECTANGULAR | CIRCULAR | ARCHED)'],
        ['explanation', 'string', 'شرح نصي للتوصية مع الأسباب'],
        ['comparisonMatrix', 'Record<Geometry, Entry>', 'مصفوفة نتائج لكل شكل'],
    ],
    col_widths=[5, 4, 7], font_size=10)

add_heading_ar(doc, 'محتويات مقارنة كل شكل (GeometryComparisonEntry)', level=3)
add_table_simple(doc,
    ['الحقل', 'الوحدة', 'الوصف'],
    [
        ['thicknessMeters', 'm', 'السماكة المطلوبة'],
        ['steelWeightTon', 'ton', 'وزن الحديد المقدّر'],
        ['concreteVolumeM3', 'm³', 'حجم الخرسانة المقدّر (لكل m طول)'],
        ['maxDynamicMomentKnM', 'kN.m', 'العزم الديناميكي الأقصى'],
        ['safetyStatus', 'enum', 'SUCCESS | WARNING | FAILURE'],
    ],
    col_widths=[5, 2, 9], font_size=10)

add_heading_ar(doc, '9.4 المعادلات الحسابية للترجيح', level=2)

add_heading_ar(doc, 'الأوزان الافتراضية', level=3)
add_table_simple(doc,
    ['المعيار', 'الوزن', 'المنطق'],
    [
        ['thicknessWeight', '35', 'السماكة الأخف أفضل (كلفة أقل)'],
        ['steelWeight', '20', 'حديد أقل = كلفة أقل'],
        ['safetyWeight', '35', 'الأمان أهم معيار'],
        ['ductilityWeight', '10', 'المطاوعة الأعلى أفضل'],
    ],
    col_widths=[4, 3, 9], font_size=10)

add_heading_ar(doc, 'درجة الأمان لكل حالة', level=3)
add_formula_box(doc, 'safetyScore(SUCCESS) = 100')
add_formula_box(doc, 'safetyScore(WARNING) = 40')
add_formula_box(doc, 'safetyScore(FAILURE) = -50')

add_heading_ar(doc, 'تطبيع السماكة والحديد (القيمة الأقل أفضل)', level=3)
add_formula_box(doc, 'normalizedScore = max(0, 100 × (1 − value/max))')

add_heading_ar(doc, 'تطبيع المطاوعة (القيمة الأعلى أفضل)', level=3)
add_formula_box(doc, 'normalizedScore = max(0, 100 × value/max)')

add_heading_ar(doc, 'الدرجة الكلية', level=3)
add_formula_box(doc, 'Score = thicknessScore×35 + steelScore×20 + safetyScore×35 + ductilityScore×10')

add_heading_ar(doc, 'حساب وزن الحديد التقريبي', level=3)
add_formula_box(doc, 'steelWeight = As × 0.00785 × thickness × 7850 / 1000  [ton]')

add_heading_ar(doc, 'حساب حجم الخرسانة التقريبي', level=3)
add_formula_box(doc, 'RECTANGULAR:  V = thickness × 10  [m³/m طول]')
add_formula_box(doc, 'CIRCULAR:     V = thickness × 8   [m³/m طول]')
add_formula_box(doc, 'ARCHED:       V = thickness × 6   [m³/m طول]')

add_heading_ar(doc, '9.5 مثال على قرار المفاضلة', level=2)
add_info_box(doc, 'مثال تطبيقي',
    [
        'إذا كانت النتائج:',
        '• مستطيل: سماكة 0.80m، أمان SUCCESS، درجة 85',
        '• دائري:  سماكة 0.65m، أمان SUCCESS، درجة 92',
        '• مقوّس:  سماكة 0.72m، أمان WARNING، درجة 60',
        '',
        'التوصية النهائية: الدائري (أعلى درجة = 92)',
        'السبب: أقل سماكة + أمان تام + مطاوعة جيدة',
    ],
    color_bg='E8F5E9', color_border='2E7D32', icon='✓')

add_page_break(doc)

print("✓ تم كتابة الفصل 9")

# ═══════════════════════════════════════════════════════════════════════
# الفصل العاشر: صفحات لوحة التحكم التفصيلية
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل العاشر: صفحات لوحة التحكم التفصيلية', level=1, color=NAVY)

add_para_ar(doc,
    'لوحة التحكم هي المساحة الرئيسية للمستخدمين بعد تسجيل الدخول. تتكوّن من 12 '
    'صفحة تفصيلية، كل منها متخصصة في جانب من جوانب خط الحساب أو في وظيفة '
    'إدارية. جميع الصفحات تتشارك تخطيطاً موحداً (RTL + Dark Theme) وقائمة '
    'جانبية (AppSidebar) للتنقل السريع. في هذا الفصل نستعرض كل صفحة على حدة '
    'مع تفصيل المدخلات والمخرجات والحسابات المرتبطة بها.',
    indent_first=True)

# ─── 10.1 الصفحة الرئيسية ───
add_heading_ar(doc, '10.1 الصفحة الرئيسية /dashboard', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'الصفحة الرئيسية هي مركز التحكم القيادي. تعرض خط الحساب الكامل (Pipeline) '
    'بتصور مرئي تفاعلي عبر مكوّن PipelineTimeline الذي يعرض 7 خطوات متسلسلة '
    'بحالاتها (انتظار، قيد التشغيل، نجاح، فشل). يتيح النموذج الرئيسي V3EngineForm '
    'إدخال كل المعطيات المطلوبة وتشغيل المحرك بضغطة زر، ثم عرض النتائج بمقارنة '
    'مرجعية آلية مع BMK-02.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_para_ar(doc, 'النموذج V3EngineForm مقسّم إلى 4 أقسام:')

add_table_simple(doc,
    ['القسم', 'الحقول', 'الوصف'],
    [
        ['السلاح', 'weaponId, impactVelocity, impactAngleDeg', 'اختيار السلاح من المكتبة + السرعة + الزاوية'],
        ['التربة', 'soilTypeCode', 'اختيار نوع التربة من 6 خيارات'],
        ['الهندسة', 'tunnelSpanShort (ap), tunnelSpanLong (bp)', 'أبعاد النفق'],
        ['التصميم', 'fcMpa, fyMpa, slabThicknessHintMm', 'خصائص المواد'],
    ],
    col_widths=[3, 6, 7], font_size=10)

add_heading_ar(doc, 'المخرجات', level=3)
add_bullet_ar(doc, 'PipelineTimeline: تصور مرئي لـ 7 خطوات بحالاتها (انتظار/تشغيل/نجاح/فشل)')
add_bullet_ar(doc, 'جدول مقارنة النتائج: القيمة المحسوبة + المرجعية + الانحراف %')
add_bullet_ar(doc, 'بطاقة قرار التحقق: SUCCESS / WARNING / FAILURE مع التفاصيل')
add_bullet_ar(doc, 'بطاقة التوصية النهائية: الشكل الموصى به (مستطيل/دائري/مقوّس) + الأسباب')
add_bullet_ar(doc, 'زر "حفظ السيناريو": يحفظ النتائج في ProjectRepository محلياً')

add_heading_ar(doc, 'الحسابات والعمليات', level=3)
add_para_ar(doc, 'عند الضغط على زر "تشغيل المحرك":')
steps_dashboard = [
    ('1. التحقق من الجلسة', 'فحص cookie sbv3-session + فك التشفير AES-256-GCM'),
    ('2. التحقق من الصلاحية', 'فحص الإذن engine:run في RBAC'),
    ('3. تشغيل المنسّق', 'استدعاء runEngine(input) من orchestrator.ts'),
    ('4. تنفيذ 4 محركات', 'الاختراق → الضغط → التصميم (3 أشكال) → المفاضلة'),
    ('5. حساب الانحراف', 'مقارنة كل قيمة مع BMK-02 المرجعية'),
    ('6. عرض النتائج', 'تحديث PipelineTimeline + جدول المقارنة + بطاقة القرار'),
    ('7. تسجيل التدقيق', 'INSERT INTO audit_logs (action=ENGINE_RUN, userId, input, output)'),
]
for title, desc in steps_dashboard:
    add_bullet_ar(doc, f'{title}: {desc}')

add_page_break(doc)

# ─── 10.2 صفحة الاختراق ───
add_heading_ar(doc, '10.2 صفحة الاختراق /dashboard/penetration', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة تفصيلية لمحرك الاختراق تعرض كل القيم الوسيطة والنهائية مع المقارنة المرجعية '
    'BMK-02. تتيح تعديل المدخلات (السلاح، السرعة، الزاوية) وإعادة التشغيل الفوري. '
    'تحتوي على 3 تبويبات: النتائج، المعادلات، والمرجعية.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_bullet_ar(doc, 'weaponId (اختيار من قائمة منسدلة بالأسلحة الأمريكية والروسية)')
add_bullet_ar(doc, 'impactVelocity (m/s) — مع شريط تمرير سريع')
add_bullet_ar(doc, 'soilTypeCode (6 خيارات مع أوصاف عربية)')
add_bullet_ar(doc, 'impactAngleDeg (0-90 درجة)')

add_heading_ar(doc, 'المخرجات', level=3)
add_table_simple(doc,
    ['القيمة', 'الوحدة', 'الوصف', 'المقارنة'],
    [
        ['x₁', 'm', 'عمق الاختراق', 'BMK-02: 3.6495 m'],
        ['h₀', 'm', 'عمق الحفرة', 'محسوب من hz × 1.18'],
        ['C_ef', 'kg', 'الشحنة الفعّالة', 'BMK-02: 334.766 kg'],
        ['λ₁', '—', 'معامل شكل الرأس', 'BMK-02: 1.1347'],
        ['λ₂', '—', 'معامل القطر', 'BMK-02: 1.2125'],
        ['n', '—', 'أُس التأثير', 'محسوب من 3.5 − Lh/D'],
        ['R_actual', 'm', 'البعد الفعلي', 'BMK-02: 7.6231 m'],
        ['h_z', 'm', 'العمق الصافي', 'BMK-02: 2.6895 m'],
        ['h̄_z', '—', 'العمق المختزل', 'BMK-02: 0.3881'],
        ['Zp', 'm/kg^1/3', 'البعد المختزل', 'BMK-02: 5.8213'],
    ],
    col_widths=[3, 2, 4, 6], font_size=10)

add_heading_ar(doc, 'الحسابات المعروضة', level=3)
add_para_ar(doc,
    'تعرض الصفحة كل معادلة من المعادلات 13-19 مع القيم العددية المُعوّضة، '
    'مما يتيح للمستخدم تتبع كل خطوة حسابية يدوياً. كما تعرض التحذيرات '
    '(ERR-PEN-01, WARN-PEN-01/02/03) في حالة وجودها، مع شرح تفصيلي لكل تحذير.')

add_page_break(doc)

# ─── 10.3 صفحة أحمال الانفجار ───
add_heading_ar(doc, '10.3 صفحة أحمال الانفجار /dashboard/blast-loads', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة تفصيلية لمحرك الضغط العصفي مع المسارين المنفصلين (سقف + جدار). '
    'تعرض تبويبين رئيسيين: تبويب السقف وتبويب الجدار، كل منهما يحتوي على '
    'جدول القيم المحسوبة + المقارنة مع STEP5_ROOF أو STEP5_WALL المرجعية.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_para_ar(doc, 'تُؤخذ القيم المقفلة من الخطوات السابقة + قيم مسارية محددة:')
add_bullet_ar(doc, 'C_ef (من محرك الاختراق — مقفل)')
add_bullet_ar(doc, 'h_pr, R_actual (من محرك الاختراق — مقفلة)')
add_bullet_ar(doc, 'Hp_cm (سماكة السقف) و Hc_cm (سماكة الجدار)')
add_bullet_ar(doc, 'Kp, Kd (معاملات الضغط والديناميكية)')
add_bullet_ar(doc, 'R_bar_b1, a0z, a1z (معاملات الاستيفاء من جدول B)')

add_heading_ar(doc, 'المخرجات لكل مسار', level=3)
add_table_simple(doc,
    ['القيمة', 'الوحدة', 'سقف', 'جدار', 'الوصف'],
    [
        ['P_max', 'kg/cm²', '4.6084', '—', 'الضغط الأقصى'],
        ['P_ekv', 'kg/cm²', '3.8158', '—', 'الضغط المكافئ'],
        ['Pct', 'kg/cm²', '1.1053', '—', 'الضغط الثابت'],
        ['Pp', 'kg/cm²', '4.9211', '3.7845', 'ضغط التصميم النهائي'],
        ['τ', 's', '0.2650', '—', 'زمن الطور الموجب'],
        ['τ_ef', 's', '0.2378', '—', 'الزمن الفعال'],
        ['ω', 'rad/s', '561.67', '—', 'التردد الدائري'],
        ['kpsi', '—', '0.90', '0.85', 'معامل المسار'],
        ['R_ekv', 'm', '6.1162', '—', 'البعد المكافئ'],
        ['R*', 'm', '2.4255', '—', 'البعد النجمي'],
        ['R_ekv > R* ?', 'boolean', 'true', '—', 'شرط الموجة'],
    ],
    col_widths=[3, 2.5, 3, 3, 4], font_size=9)

add_heading_ar(doc, 'الحسابات المعروضة', level=3)
add_para_ar(doc,
    'تعرض الصفحة كل معادلة من المعادلات (1) إلى (8) + المعادلات الإضافية لـ blast-loads.ts '
    'مع القيم العددية. كما تعرض نسبة الانحراف عن المرجع مع رمز لوني: أخضر (< 2%)، '
    'أصفر (2-5%)، أحمر (> 5%). القيم التي تنحرف أكثر من 5% تُستبدل آلياً بالقيم المرجعية المقفلة.')

add_page_break(doc)

# ─── 10.4 صفحة التصميم الإنشائي ───
add_heading_ar(doc, '10.4 صفحة التصميم الإنشائي /dashboard/structural-design', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة التصميم الإنشائي تعرض نتائج محرك structural-concrete-core.ts للأشكال '
    'الثلاثة (مستطيل، دائري، مقوّس) بشكل متوازٍ. تتضمن رسماً بيانياً SVG لمقطع '
    'كل شكل مع الأبعاد المحسوبة (السماكة، التسليح). كما تعرض فحوصات التحقق '
    '(اللامركزية، القص الثاقب، نسبة التسليح) برمز لوني.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_bullet_ar(doc, 'pDesignMpa (من محرك الضغط — مقفل)')
add_bullet_ar(doc, 'geometryType (3 قيم: RECTANGULAR, CIRCULAR, ARCHED)')
add_bullet_ar(doc, 'tunnelSpanShort (ap), tunnelSpanLong (bp)')
add_bullet_ar(doc, 'fcMpa, fyMpa (خصائص المواد)')

add_heading_ar(doc, 'المخرجات لكل شكل', level=3)
add_table_simple(doc,
    ['القيمة', 'الوحدة', 'الوصف'],
    [
        ['requiredThicknessMeters', 'm', 'السماكة المطلوبة'],
        ['requiredSteelAreaCm2PerMeter', 'cm²/m', 'مساحة التسليح المطلوبة'],
        ['ductilityRatio', '—', 'نسبة المطاوعة = fy/fc'],
        ['eccentricityRatio', '—', 'نسبة اللامركزية e/h'],
        ['punchingShearRatio', '—', 'إجهاد القص الفعلي/المسموح'],
        ['reinforcementRatio', '—', 'نسبة التسليح الفعلية'],
        ['status', 'enum', 'SUCCESS | WARNING | FAILURE'],
        ['failures', 'string[]', 'قائمة الإخفاقات التفصيلية'],
    ],
    col_widths=[5, 2, 9], font_size=10)

add_heading_ar(doc, 'الحسابات المعروضة', level=3)
add_para_ar(doc,
    'تعرض الصفحة: (1) تطبيق DIF على fc و fy، (2) حساب العزم الديناميكي '
    'M = P × ap² / 8، (3) حساب السماكة المطلوبة h_required، (4) حساب مساحة '
    'التسليح As، (5) فحص اللامركزية e ≤ h/6، (6) فحص القص الثاقب v ≤ v_cd. '
    'كل فحص يُعرض مع رمز ✓ (نجاح) أو ✗ (فشل) أو ⚠ (تحذير).')

add_page_break(doc)

# ─── 10.5 صفحة التسليح ───
add_heading_ar(doc, '10.5 صفحة التسليح /dashboard/rebar-design', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة تصميم التسليح تعرض تفاصيل محرك rebar.ts لكلا المسارين (سقف + جدار). '
    'تتضمن جدول اختيار القضبان الأمثل، فحوصات الكود (ρ_min, ρ_max, ξ_max)، '
    'و رسماً تخطيطياً لتوزيع القضبان في المقطع.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_bullet_ar(doc, 'Mp (العزم البلاستيكي — من محرك التصميم)')
add_bullet_ar(doc, 'h0 (العمق الفعلي — من محرك التصميم)')
add_bullet_ar(doc, 'Rsd, Rbd (المقاومات الديناميكية)')
add_bullet_ar(doc, 'path (roof | wall)')

add_heading_ar(doc, 'المخرجات', level=3)
add_table_simple(doc,
    ['القيمة', 'الوحدة', 'الوصف'],
    [
        ['As_required', 'cm²/m', 'مساحة التسليح المطلوبة (محسوبة)'],
        ['As_provided', 'cm²/m', 'مساحة التسليح المقدمة (مختارة)'],
        ['barCount', '—', 'عدد القضبان'],
        ['barDiameter', 'mm', 'قطر القضيب'],
        ['ρ (rho)', '—', 'نسبة التسليح الفعلية'],
        ['ξ (xi)', '—', 'العمق النسبي لمنطقة الضغط'],
        ['αm', '—', 'معامل العزم النسبي'],
        ['status', 'enum', 'OK | WARNING | FAILURE'],
    ],
    col_widths=[4, 2, 9], font_size=10)

add_heading_ar(doc, 'الفحوصات المعروضة', level=3)
add_table_simple(doc,
    ['الفحص', 'الشرط', 'الكود', 'الحالة'],
    [
        ['minReinforcementOk', 'ρ ≥ 0.003', 'الكود السوري 2024', 'إذا فشل → WARNING'],
        ['maxReinforcementOk', 'ρ ≤ 0.04', 'UFC 3-340-02', 'إذا فشل → FAILURE'],
        ['xiConditionOk', 'ξ ≤ 0.65', 'UFC 3-340-02', 'إذا فشل → FAILURE'],
    ],
    col_widths=[4, 3, 4, 5], font_size=10)

add_page_break(doc)

# ─── 10.6 صفحة المفاضلة ───
add_heading_ar(doc, '10.6 صفحة المفاضلة /dashboard/comparison', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة المفاضلة الهندسية تعرض نتائج محرك geometry-comparator.ts بشكل بصري. '
    'تتضمن جدول مقارنة شامل للأشكال الثلاثة (مستطيل، دائري، مقوّس) مع 5 معايير '
    'لكل شكل، وبطاقة التوصية النهائية prominently displayed.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_para_ar(doc, 'تُستقبل من API /api/comparison الذي يستدعي compareGeometries:')

add_heading_ar(doc, 'المخرجات', level=3)
add_table_simple(doc,
    ['الشكل', 'السماكة (m)', 'وزن الحديد (ton)', 'حجم الخرسانة (m³)', 'الأمان', 'الدرجة'],
    [
        ['مستطيل', '0.80', '0.62', '8.0', 'SUCCESS', '85'],
        ['دائري', '0.65', '0.51', '5.2', 'SUCCESS', '92 ⭐'],
        ['مقوّس', '0.72', '0.56', '4.3', 'WARNING', '60'],
    ],
    col_widths=[3, 3, 3, 3.5, 2.5, 2], font_size=10)

add_para_ar(doc, 'بطاقة التوصية النهائية تعرض: الشكل الموصى به (دائري)، '
                  'الدرجة الكلية (92)، شرح التوصية، ومقارنة سريعة مع الخيارات الأخرى.')

add_page_break(doc)

# ─── 10.7 صفحة خط الأساس ───
add_heading_ar(doc, '10.7 صفحة خط الأساس /dashboard/benchmark', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة خط الأساس تتيح تشغيل اختبارات الانحدار (Regression Tests) على المحركات '
    'للتحقق من دقتها مقابل الحالات المرجعية الثلاثة: BMK-01، BMK-02، BMK-03. '
    'كل حالة لها مدخلات محددة وقيم متوقعة مع تسامح، وتُعتبر المحركات صحيحة '
    'إذا كانت كل الانحرافات ضمن التسامح المسموح.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_bullet_ar(doc, 'اختيار الحالة المرجعية (BMK-01/02/03)')
add_bullet_ar(doc, 'معاملات التسامح (افتراضي 2%)')

add_heading_ar(doc, 'المخرجات', level=3)
add_table_simple(doc,
    ['الحقل', 'الوصف'],
    [
        ['benchmarkId', 'معرف الحالة (BMK-01/02/03)'],
        ['timestamp', 'طابع زمني للتشغيل'],
        ['results[]', 'مصفوفة نتائج لكل قيمة متوقعة'],
        ['overallPassed', 'boolean — نجاح شامل'],
        ['summary.total', 'عدد القيم المختبرة'],
        ['summary.passed', 'عدد القيم الناجحة'],
        ['summary.failed', 'عدد القيم الفاشلة'],
        ['summary.maxDeviationPercent', 'أقصى انحراف مئوي'],
    ],
    col_widths=[5, 11], font_size=10)

add_heading_ar(doc, 'الحالات المرجعية الثلاث', level=3)
add_table_simple(doc,
    ['الحالة', 'السلاح', 'التربة', 'السرعة (m/s)', 'الأولوية'],
    [
        ['BMK-01', 'MK82', 'SOFT_SOIL', '250', 'حرجة (1)'],
        ['BMK-02', 'MK83', 'MEDIUM_SOIL', '350', 'حرجة (1) — الحالة الرئيسية'],
        ['BMK-03', 'MK84', 'HARD_ROCK', '450', 'مهمة (2)'],
    ],
    col_widths=[3, 3, 4, 3, 5], font_size=10)

add_page_break(doc)

# ─── 10.8 صفحة المتغيرات الموحدة ───
add_heading_ar(doc, '10.8 صفحة المتغيرات الموحدة /dashboard/variables', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة المتغيرات الموحدة تعرض جدولاً تفصيلياً بكل المتغيرات الـ 42 في النظام، '
    'مصنّفة حسب الفئة (input/lookup/computed/locked/output) والمسار (سقف/جدار/مشترك) '
    'والمحرك المسؤول. تتيح الصفحة فلترة المتغيرات وبحثها، وعرض تبعيات كل متغير.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_bullet_ar(doc, 'فلاتر: الفئة، المسار، المحرك، حالة القفل')
add_bullet_ar(doc, 'بحث نصي في الاسم أو الرمز أو الوصف')

add_heading_ar(doc, 'المخرجات', level=3)
add_table_simple(doc,
    ['العمود', 'الوصف'],
    [
        ['الرمز', 'الرمز الرياضي للمتغير (مثل P, d, C_ef)'],
        ['الاسم', 'الاسم العربي المختصر'],
        ['الوحدة', 'وحدة القياس'],
        ['الفئة', 'input | lookup | computed | locked | output'],
        ['المسار', 'roof | wall | shared'],
        ['المحرك', 'penetration | blast | structural'],
        ['مقفل؟', 'boolean — هل القيمة مقفلة (للقراءة فقط)'],
        ['يستند إلى', 'قائمة المتغيرات التي يعتمد عليها'],
        ['المعادلة', 'المعادلة الحسابية (إن وجدت)'],
    ],
    col_widths=[3, 13], font_size=10)

add_heading_ar(doc, 'إحصائيات التصنيف', level=3)
add_table_simple(doc,
    ['الفئة', 'العدد التقريبي', 'الوصف'],
    [
        ['input', '8', 'مدخلات المستخدم المباشرة'],
        ['lookup', '12', 'قيم من جداول المرجعية'],
        ['computed', '10', 'محسوبة داخل محرك واحد'],
        ['locked', '7', 'مقفلة بين المحركات'],
        ['output', '5', 'المخرجات النهائية'],
    ],
    col_widths=[3, 3, 10], font_size=10)

add_page_break(doc)

# ─── 10.9 صفحة مصفوفة التتبع ───
add_heading_ar(doc, '10.9 صفحة مصفوفة التتبع /dashboard/rtm', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة مصفوفة التتبع (Requirements Traceability Matrix) هي أداة جودة تربط كل '
    'متطلب تصميمي بالنتيجة المحسوبة وبالحالة المرجعية. تتأكد من أن كل متطلب في '
    'الأطروحة قد تم تنفيذه في المنصة وأن النتيجة تتطابق مع المرجع ضمن التسامح.',
    indent_first=True)

add_heading_ar(doc, 'المكونات الرئيسية', level=3)
add_bullet_ar(doc, 'VariableTraceabilityMatrix: جدول يربط المتغيرات بالمعادلات والمراجع')
add_bullet_ar(doc, 'AuditTrailEnhanced: سجل تدقيق محسّن لكل عملية حسابية')
add_bullet_ar(doc, 'DefectLog: سجل العيوب والانحرافات عن المرجع')
add_bullet_ar(doc, 'RtmDashboardController: لوحة قيادة للتحكم بالعرض والفلاتر')

add_heading_ar(doc, 'المدخلات والمخرجات', level=3)
add_table_simple(doc,
    ['العنصر', 'المدخلات', 'المخرجات'],
    [
        ['Traceability Matrix', 'قائمة المتغيرات + المعادلات', 'جدول تتبع كامل'],
        ['Audit Trail', 'userId + sessionId', 'سجل زمني لكل العمليات'],
        ['Defect Log', 'نتائج BMK', 'قائمة الانحرافات مع التشخيص'],
    ],
    col_widths=[4, 5, 7], font_size=10)

add_page_break(doc)

# ─── 10.10 صفحة التقارير ───
add_heading_ar(doc, '10.10 صفحة التقارير /dashboard/reports', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة التقارير تتيح للمستخدم إنشاء تقارير PDF احترافية لنتائج المحرك. '
    'كل تقرير يتضمن: ملخص تنفيذي، المدخلات، النتائج الوسيطة، النتائج النهائية، '
    'مصفوفة التحقق، والتوصية. تتطلب إذن reports:export للتصدير.',
    indent_first=True)

add_heading_ar(doc, 'المدخلات', level=3)
add_bullet_ar(doc, 'scenarioId (معرّف السيناريو المحفوظ)')
add_bullet_ar(doc, 'reportFormat (PDF | Excel)')
add_bullet_ar(doc, 'includeSections (قائمة الأقسام المطلوبة)')

add_heading_ar(doc, 'المخرجات', level=3)
add_bullet_ar(doc, 'ملف PDF أو Excel قابل للتنزيل')
add_bullet_ar(doc, 'تسجيل في audit_logs (action=REPORT_EXPORTED)')
add_bullet_ar(doc, 'تحديث counter في جدول reports_generated')

add_page_break(doc)

# ─── 10.11 صفحة الإعدادات ───
add_heading_ar(doc, '10.11 صفحة الإعدادات /dashboard/settings', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة الإعدادات تتيح للمستخدم تخصيص بعض جوانب المنصة: أوزان المفاضلة، '
    'حالة الـ PWA، المزامنة، السمات (Themes)، وإعدادات الأمان. تتضمن أيضاً '
    'معلومات الحساب وإمكانية تغيير كلمة المرور.',
    indent_first=True)

add_heading_ar(doc, 'الأقسام', level=3)
add_table_simple(doc,
    ['القسم', 'الإعدادات المتاحة'],
    [
        ['الحساب', 'تغيير كلمة المرور، الاسم، المؤسسة'],
        ['الأوزان', 'thicknessWeight, steelWeight, safetyWeight, ductilityWeight'],
        ['المزامنة', 'تفعيل/تعطيل، تكرار المزامنة، سياسة حل التعارض'],
        ['العرض', 'السمة (فاتح/داكن)، الكثافة، اللغة (عربي/إنجليزي)'],
        ['الأمان', 'إدارة الجلسات، الإشعارات، تسجيل الخروج من كل الأجهزة'],
        ['PWA', 'تثبيت التطبيق، تحديث Service Worker، حالة الكاش'],
    ],
    col_widths=[3, 13], font_size=10)

add_page_break(doc)

# ─── 10.12 صفحة عن المنصة ───
add_heading_ar(doc, '10.12 صفحة عن المنصة /dashboard/about', level=2)

add_heading_ar(doc, 'الهدف', level=3)
add_para_ar(doc,
    'صفحة معلوماتية تعرض: المراجع العلمية، الإصدار، السجل الزمني للتطوير، '
    'فريق العمل، الترخيص، وروابط التواصل. كما تتضمن قائمة بالأخطاء المعروفة '
    'وخطوات استكشافها وإصلاحها.',
    indent_first=True)

add_heading_ar(doc, 'المحتويات', level=3)
add_bullet_ar(doc, 'المرجعية العلمية: الأطروحة + ملفات Excel I-1 إلى I-9')
add_bullet_ar(doc, 'الكود الحاكم: الكود السوري 2024 + UFC 3-340-02')
add_bullet_ar(doc, 'الإصدار: V3.0 — تاريخ الإصدار 2026-06')
add_bullet_ar(doc, 'التراخيص: ملكية سيادية — استخدام محدود')
add_bullet_ar(doc, 'تقنيات المنصة: Next.js 16 + TypeScript + Tailwind 4 + shadcn/ui')
add_bullet_ar(doc, 'قاعدة البيانات: Supabase PostgreSQL عبر REST API')
add_bullet_ar(doc, 'الاستضافة: Netlify مع Serverless Functions')

add_page_break(doc)

print("✓ تم كتابة الفصل 10")

# ═══════════════════════════════════════════════════════════════════════
# الفصل الحادي عشر: المعادلات الحسابية المرجعية
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الحادي عشر: المعادلات الحسابية المرجعية', level=1, color=NAVY)

add_heading_ar(doc, '11.1 فهرس المعادلات الكامل', level=2)
add_para_ar(doc,
    'هذا الفصل يجمع كل المعادلات المستخدمة في المنصة بترتيبها الرقمي المرجعي. '
    'كل معادلة لها رقم مرجعي يربطها بالأطروحة الأصلية وملفات Excel. هذا الفهرس '
    'هو المرجع النهائي لفهم الحسابات والتحقق منها.')

add_heading_ar(doc, 'محرك الضغط العصفي (المعادلات 1-12)', level=2)

add_table_simple(doc,
    ['رقم', 'الاسم', 'الصيغة', 'الوحدة'],
    [
        ['1', 'الضغط الزائد (سادوفسكي)', 'ΔP = 0.1∛C/R + 0.43∛C²/R² + 1.4C/R³', 'MPa'],
        ['2', 'الضغط المنعكس', 'Pr = 2 × Kp × σ_max', 'MPa'],
        ['3', 'الإجهاد الأقصى في التربة', 'σ_max = A × Z^(-n1)', 'MPa'],
        ['4', 'زمن الطور الموجب', 'τ⁺ = 1.7×10⁻³ × ∛(√C) × √R', 's'],
        ['5', 'الاندثار (Impulse)', 'I = 6.3 × ∛(C²) / R', '—'],
        ['6', 'البعد المختزل', 'Z = R / ∛C', 'm/kg^1/3'],
        ['7', 'البعد الحرج', 'R_critical = 1.1 × ∛C', 'm'],
        ['8', 'الزمن الفعال', 'τ_eff = τ⁺ × f(ΔPmax)', 's'],
        ['9', 'معامل الاستيفاء I-9', 'f = 0.0008ΔP² − 0.0384ΔP + 1.0013', '—'],
        ['10', 'الضغط التصميمي', 'P_design = P_static + Kd × P_max', 'MPa'],
        ['11', 'الضغط الساكن', 'P_static = ceilingDepth × 2000/10000 × 0.0981', 'MPa'],
        ['12', 'شرط الديناميكية', 'τ_eff ≥ 0.2 × π / ω', 'boolean'],
    ],
    col_widths=[1, 4, 8, 3], font_size=9)

add_heading_ar(doc, 'محرك الاختراق (المعادلات 13-19)', level=2)

add_table_simple(doc,
    ['رقم', 'الاسم', 'الصيغة', 'الوحدة'],
    [
        ['13', 'عمق الاختراق', 'x₁ = λ₁ × λ₂ × Kpr × (P/d²) × V × cos(α)', 'm'],
        ['14', 'معامل شكل الرأس', 'λ₁ = 0.5 + 0.4 × (Lh/D)^(2/3)', '—'],
        ['15', 'معامل القطر', 'λ₂ = 2.8 × d^(1/3) − 1.3 × d^(1/2)', '—'],
        ['16', 'أُس التأثير', 'n = 3.5 − (Lh/D)', '—'],
        ['17', 'زاوية الاختراق (خارقة)', 'τ = 0.5 × lₖ × cos((α + n×α)/2)', 'm'],
        ['18', 'زاوية الاختراق (انفجارية)', 'τ = 0.5 × dₖ', 'm'],
        ['19', 'الشحنة الفعّالة', 'C_ef = 0.95 × K₁ × C', 'kg'],
    ],
    col_widths=[1, 4, 8, 3], font_size=9)

add_heading_ar(doc, 'محرك التصميم الإنشائي', level=2)

add_table_simple(doc,
    ['الاسم', 'الصيغة', 'الوحدة'],
    [
        ['معامل التضخيم الخرساني', 'fcd = fc × 1.25', 'MPa'],
        ['معامل التضخيم للحديد', 'fsd = fy × 1.20', 'MPa'],
        ['العمق الفعلي', 'd_eff = h − 50mm', 'mm'],
        ['العزم الديناميكي', 'M = P × ap² / 8', 'kN.m'],
        ['السماكة المطلوبة', 'h = √(M×10⁵ / (Rb × b × 0.3)) + cover', 'mm'],
        ['مساحة التسليح', 'As = (M×10⁶) / (fsd × 0.875 × d_eff)', 'mm²'],
        ['نسبة التسليح', 'ρ = As / (d_eff × 100)', '—'],
        ['حد اللامركزية', 'e_limit = h / 6', 'mm'],
        ['القص الثاقب المسموح', 'v_cd = 0.25 × √(fcd)', 'MPa'],
        ['المحيط الحرج', 'b₀ = 2(b_col + d_eff) + 2(h_col + d_eff)', 'mm'],
    ],
    col_widths=[5, 9, 2], font_size=10)

add_heading_ar(doc, 'محرك التسليح', level=2)

add_table_simple(doc,
    ['الاسم', 'الصيغة', 'الوحدة'],
    [
        ['مساحة التسليح المطلوبة', 'As = Mp / (Rsd × z), z = h0 − 0.5x', 'cm²/m'],
        ['عمق منطقة الضغط', 'x = Rsd × As / (Rbd × b)', 'cm'],
        ['العمق النسبي', 'ξ = Rsd × As / (Rbd × b × h0)', '—'],
        ['معامل العزم النسبي', 'αm = ξ × (1 − 0.5ξ)', '—'],
        ['نسبة التسليح', 'ρ = As / (b × h0)', '—'],
        ['مقاومة القص الديناميكية', 'Rsd = RsH × 1.2 × 1.25', 'kg/cm²'],
        ['مقاومة الانحناء الديناميكية', 'Rbd = RbH × 1.25 × 1.25 / 10', 'kg/cm²'],
        ['نسبة المطاوعة الإنشائية', 'μ = (RsH/RbH) × ξ × (1 − 0.5ξ)', '—'],
    ],
    col_widths=[5, 9, 2], font_size=10)

add_heading_ar(doc, 'محرك المفاضلة', level=2)

add_table_simple(doc,
    ['الاسم', 'الصيغة', 'الوحدة'],
    [
        ['درجة الأمان', 'SUCCESS=100, WARNING=40, FAILURE=-50', 'نقطة'],
        ['تطبيع عكسي (أقل أفضل)', 'score = 100 × (1 − value/max)', 'نقطة'],
        ['تطبيع مباشر (أعلى أفضل)', 'score = 100 × value/max', 'نقطة'],
        ['الدرجة الكلية', 'Score = 35t + 20s + 35a + 10d', 'نقطة'],
        ['وزن الحديد التقريبي', 'steel = As × 0.00785 × h × 7.85', 'ton'],
        ['حجم الخرسانة', 'V_rect=h×10, V_circ=h×8, V_arch=h×6', 'm³/m'],
    ],
    col_widths=[5, 9, 2], font_size=10)

add_page_break(doc)

print("✓ تم كتابة الفصل 11")

# ═══════════════════════════════════════════════════════════════════════
# الفصل الثاني عشر: جدول المتغيرات الموحدة
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الثاني عشر: جدول المتغيرات الموحدة (42 متغير)', level=1, color=NAVY)

add_para_ar(doc,
    'هذا الجدول هو المصدر المرجعي لكل متغير في المنصة. كل متغير له رمز، وحدة، '
    'فئة، مسار، محرك مسؤول، وقائمة التبعيات. القيم المقفلة (locked) تنتج من محرك '
    'وتُستهلك في المحرك التالي دون إمكانية إعادة الكتابة، مما يضمن عدم الانحراف '
    'عن المرجع.',
    indent_first=True)

add_heading_ar(doc, '12.1 المتغيرات المدخلة (input)', level=2)
add_table_simple(doc,
    ['الرمز', 'الوحدة', 'الوصف', 'المسار'],
    [
        ['P', 'kg', 'وزن القنبلة', 'مشترك'],
        ['d', 'm', 'قطر القنبلة', 'مشترك'],
        ['V', 'm/s', 'سرعة الاصطدام', 'مشترك'],
        ['α', 'درجة', 'زاوية الاصطدام', 'مشترك'],
        ['C', 'kg', 'وزن الشحنة المتفجرة', 'مشترك'],
        ['ap', 'm', 'البحر القصير للنفق', 'مشترك'],
        ['bp', 'm', 'البحر الطويل للنفق', 'مشترك'],
        ['Z', 'm', 'عمق الدفن', 'مشترك'],
    ],
    col_widths=[3, 2, 8, 3], font_size=10)

add_heading_ar(doc, '12.2 المتغيرات المرجعية (lookup)', level=2)
add_table_simple(doc,
    ['الرمز', 'الوحدة', 'الوصف', 'المصدر'],
    [
        ['K₁', '—', 'معامل التحويل إلى TNT', 'جدول I-5'],
        ['kpr', '—', 'معامل اختراق التربة', 'جدول I-1'],
        ['A', '—', 'معامل الإجهاد', 'جدول I-3'],
        ['n1', '—', 'أُس الإجهاد', 'جدول I-3'],
        ['a0z', '—', 'معامل السرعة الثابت', 'جدول I-2'],
        ['a1z', '—', 'معامل السرعة المتغير', 'جدول I-2'],
        ['γ_b', 'kg/m³', 'كثافة الخرسانة', 'ثابت فيزيائي'],
        ['γ_g', 'kg/m³', 'كثافة التربة', 'ثابت فيزيائي'],
        ['Ea', 'kg/cm²', 'معامل المرونة', 'جدول I-7'],
        ['RbH', 'kg/cm²', 'مقاومة الانحناء', 'جدول I-6'],
        ['RsH', 'kg/cm²', 'إجهاد الخضوع للحديد', 'جدول I-7'],
        ['Kpod', '—', 'معامل التأسيس', 'جدول I-4'],
    ],
    col_widths=[3, 2, 6, 5], font_size=10)

add_heading_ar(doc, '12.3 المتغيرات المحسوبة (computed)', level=2)
add_table_simple(doc,
    ['الرمز', 'الوحدة', 'الوصف', 'المحرك'],
    [
        ['λ₁', '—', 'معامل شكل الرأس', 'الاختراق'],
        ['λ₂', '—', 'معامل القطر', 'الاختراق'],
        ['n', '—', 'أُس التأثير', 'الاختراق'],
        ['τ', 'm', 'معامل زاوية الاختراق', 'الاختراق'],
        ['Z (مختزل)', 'm/kg^1/3', 'البعد المختزل', 'الضغط'],
        ['σ_max', 'MPa', 'الإجهاد الأقصى', 'الضغط'],
        ['Pr', 'MPa', 'الضغط المنعكس', 'الضغط'],
        ['τ⁺', 's', 'زمن الطور الموجب', 'الضغط'],
        ['f(ΔP)', '—', 'دالة الاستيفاء', 'الضغط'],
        ['M', 'kN.m', 'العزم الديناميكي', 'التصميم'],
    ],
    col_widths=[3, 2, 7, 4], font_size=10)

add_heading_ar(doc, '12.4 المتغيرات المقفلة (locked)', level=2)
add_table_simple(doc,
    ['الرمز', 'الوحدة', 'الوصف', 'المنتج من'],
    [
        ['C_ef', 'kg', 'الشحنة الفعّالة', 'الاختراق → الضغط'],
        ['h_pr', 'm', 'عمق الاختراق', 'الاختراق → الضغط'],
        ['R_actual', 'm', 'البعد الفعلي', 'الاختراق → الضغط'],
        ['ht', 'cm', 'السماكة الكلية المقفلة', 'الخطوة 4 → التصميم'],
        ['Bt', 'm', 'البحر المكافئ', 'الخطوة 4 → التصميم'],
        ['Hpc', 'm', 'عمق مركز الانفجار', 'الخطوة 4 → التصميم'],
        ['Pp', 'kg/cm²', 'ضغط التصميم النهائي', 'الضغط → التصميم'],
    ],
    col_widths=[3, 2, 6, 5], font_size=10)

add_heading_ar(doc, '12.5 المتغيرات الناتجة (output)', level=2)
add_table_simple(doc,
    ['الرمز', 'الوحدة', 'الوصف', 'المحرك المنتج'],
    [
        ['h_required', 'm', 'السماكة المطلوبة', 'التصميم'],
        ['As', 'cm²/m', 'مساحة التسليح', 'التسليح'],
        ['ρ', '—', 'نسبة التسليح', 'التسليح'],
        ['ξ', '—', 'العمق النسبي', 'التسليح'],
        ['status', 'enum', 'حالة القرار النهائي', 'المفاضلة'],
    ],
    col_widths=[3, 2, 6, 5], font_size=10)

add_page_break(doc)

print("✓ تم كتابة الفصل 12")

# ═══════════════════════════════════════════════════════════════════════
# الفصل الثالث عشر: خلاصة وتوصيات
# ═══════════════════════════════════════════════════════════════════════

add_heading_ar(doc, 'الفصل الثالث عشر: خلاصة وتوصيات', level=1, color=NAVY)

add_heading_ar(doc, '13.1 خلاصة المنصة', level=2)
add_para_ar(doc,
    'منصة التدقيق الديناميكي الموحد V3.0 هي نظام متكامل يجمع بين 5 محركات حسابية '
    'متخصصة (الاختراق، الضغط، التصميم الإنشائي، التسليح، المفاضلة) في خط أنابيب '
    'موحد يبدأ من مدخلات المستخدم وينتهي بتوصية هندسية نهائية بشأن الشكل الإنشائي '
    'الأمثل. كل محرك مبني كدالة نقية (Pure Function) قابلة للاختبار المستقل، مع '
    'عقد واضح للمدخلات والمخرجات.',
    indent_first=True)

add_para_ar(doc,
    'تتميز المنصة بنموذج متغيرات موحّد يصنّف كل قيمة إلى 5 فئات (input/lookup/'
    'computed/locked/output) و3 مسارات (سقف/جدار/مشترك)، مما يمنع التضارب بين '
    'المحركات ويضمن قابلية التتبع. القيم المقفلة (locked) لا تُعاد كتابتها بين '
    'المحركات، وهو ضمان أساسي لمنع الانحراف التراكمي عن الحالة المرجعية BMK-02.',
    indent_first=True)

add_heading_ar(doc, '13.2 نقاط القوة', level=2)
strengths = [
    ('الدقة العلمية', 'كل معادلة مستخرجة من أطروحة مرجعية ومُنَضَّدة بملفات Excel، مع اختبارات انحدار آلية (BMK-01/02/03).'),
    ('الشفافية', 'كل خطوة حسابية موثقة ومعروضة في الواجهة، مع نسبة الانحراف عن المرجع.'),
    ('الأمان', 'تشفير AES-256-GCM للجلسات + bcrypt لكلمات المرور + RBAC بـ 7 أذونات تفصيلية.'),
    ('الحوكمة', 'حالة PENDING افتراضية + سجل تدقيق كامل + لوحة مدير سيادية.'),
    ('الأداء', 'Serverless على Netlify + Supabase REST API + IndexedDB للأوفلاين.'),
    ('التدويل', 'واجهة عربية كاملة RTL + دعم ثنائي اللغة محتمل.'),
]
for title, desc in strengths:
    add_bullet_ar(doc, f'{title}: {desc}')

add_heading_ar(doc, '13.3 التوصيات للمستخدم', level=2)
recommendations = [
    'ابدأ دائماً بمراجعة القيم المرجعية BMK-02 لضمان أن المحرك يعمل ضمن التسامح المسموح.',
    'استخدم صفحة المتغيرات الموحدة لفهم تدفق البيانات قبل تشغيل المحرك.',
    'فعّل صفحة مصفوفة التتبع للتأكد من سلامة كل متطلب تصميمي.',
    'صدّر التقارير PDF بعد كل تشغيل مهم للأرشفة والمراجعة.',
    'إذا ظهر فشل (FAILURE)، راجع صفحة المعادلات لفهم السبب الجذري قبل التعديل.',
    'احفظ السيناريوهات الناجحة في ProjectRepository لإعادة الاستخدام لاحقاً.',
]
for rec in recommendations:
    add_bullet_ar(doc, rec)

add_heading_ar(doc, '13.4 التطوير المستقبلي', level=2)
future = [
    'إضافة المزيد من الأسلحة الروسية الحديثة إلى المكتبة (مثل KAB-500).',
    'تطوير محرك التشقق (spalling) مع معاملات Kot الكاملة من جدول I-1.',
    'إضافة دعم الموجات الزلزالية للأنفاق العميقة (> 30m).',
    'تطوير واجهة ثلاثية الأبعاد (3D) لعرض المقطع والقضبان.',
    'إضافة خوارزمية تحسين (Optimization) لاختيار أفضل معاملات التصميم آلياً.',
    'تكامل مع BIM (Revit/IFC) لتصدير التصاميم مباشرة.',
]
for f in future:
    add_bullet_ar(doc, f)

add_heading_ar(doc, '13.5 مراجع التواصل', level=2)
add_para_ar(doc,
    'للاستفسارات الفنية والتقارير عن الأخطاء وطلبات التحديث، يرجى التواصل مع '
    'فريق الدعم الفني للمنصة عبر القنوات الرسمية المعتمدة. جميع التحديثات '
    'والإصدارات الجديدة تُعلن في صفحة "عن المنصة" ضمن لوحة التحكم.',
    indent_first=True)

# خاتمة
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━')
run.font.size = Pt(12)
run.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run('نهاية الوثيقة')
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = NAVY
set_rtl_run(run)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('منصة التدقيق الديناميكي الموحد V3.0')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = GRAY
set_rtl_run(run)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

print("✓ تم كتابة الفصل 13 (الخلاصة)")

# ═══════════════════════════════════════════════════════════════════════
# حفظ الوثيقة
# ═══════════════════════════════════════════════════════════════════════

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)

file_size = os.path.getsize(OUTPUT_PATH) / 1024
print(f"\n{'='*60}")
print(f"✅ تم إنشاء الوثيقة بنجاح")
print(f"📁 المسار: {OUTPUT_PATH}")
print(f"📏 الحجم: {file_size:.1f} KB")
print(f"{'='*60}")





